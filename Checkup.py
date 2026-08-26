#checkup


import streamlit as st
import pandas as pd
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
import base64
import os

# NEW: Oracle database integration
from oracle_functions import *
from oracle_ui import *


# CRITICAL: Import analyze_report from analysis_utils
from analysis_utils import analyze_report, extract_db_info, standardize_dataframes


# ==================== SECTION OBSERVATIONS AND RECOMMENDATIONS ====================
# Predefined observations, recommendations, notes, and commands for each section
SECTION_METADATA = {
    "TABLES DEGREE": {
        "observation": "Tables were found where the degree of parallelism was set to a value other than the default (1).",
        "recommendation": "Keep the degree of all tables set to 1 (default value) unless parallelism is explicitly required for specific workloads.",
        "note": "",
        "command": ""
    },
    "INDEX DEGREE": {
        "observation": "Indexes were found where the degree of parallelism was set to a value other than the default (1).",
        "recommendation": "Keep the degree of all indexes set to 1 (default value) unless parallelism is explicitly required for specific workloads.",
        "note": "",
        "command": ""
    },
    "AUD$ TABLES": {
        "observation": "AUD$ table space usage detected.",
        "recommendation": "If the AUD$ table becomes large, move it to a dedicated tablespace and purge old audit records to control size and maintain performance.",
        "note": "",
        "command": ""
    },
    "TABLESPACE INFORMATION": {
        "observation": "Tablespace usage monitored.",
        "recommendation": "Monitor the size of tablespace to ensure it doesn't cross threshold values.",
        "note": "Tablespaces approaching or exceeding threshold should be expanded or cleaned up.",
        "command": ""
    },
    "Oracle Auto-Jobs": {
        "observation": "Oracle internal jobs are running daily in background.",
        "recommendation": "Disable unnecessary Oracle auto-jobs to reduce background load.",
        "note": "Need to disable jobs on immediate basis if they conflict with business jobs.",
        "command": ""
    },
    "ORACLE AUTO-JOBS": {  # Case variation
        "observation": "Oracle internal jobs are running daily in background.",
        "recommendation": "Disable unnecessary Oracle auto-jobs to reduce background load.",
        "note": "Need to disable jobs on immediate basis if they conflict with business jobs.",
        "command": ""
    }
}


# Sections that should ONLY be shown when they have data (not just metadata)
SECTIONS_REQUIRE_DATA = ["TABLES DEGREE", "INDEX DEGREE"]

def should_show_section_without_data(section_name):
    """Check if a section should be shown even without data"""
    # Normalize section name for comparison
    section_upper = section_name.upper().strip()
    
    # Check if this section requires data
    for required_data_section in SECTIONS_REQUIRE_DATA:
        if section_upper == required_data_section.upper():
            return False  # Don't show without data
    
    return True  # Can show without data if has metadata

def get_section_metadata(section_name):
    """Get metadata for a section with case-insensitive matching"""
    # Try exact match first
    if section_name in SECTION_METADATA:
        return SECTION_METADATA[section_name]
    
    # Try case-insensitive match
    section_upper = section_name.upper()
    for key, value in SECTION_METADATA.items():
        if key.upper() == section_upper:
            return value
    
    return None


# ---------------- Helper Functions ----------------

def normalize_section_name(name):
    """Normalize section names for case-insensitive matching"""
    if not name:
        return ""
    return name.strip().upper().replace('  ', ' ')


def find_section_in_results(section_name, results):
    """
    Find section in results with case-insensitive matching
    Returns: (found_key, dataframe) or (None, None)
    """
    if not section_name or not results:
        return None, None
    
    normalized_target = normalize_section_name(section_name)
    
    # First try exact match
    if section_name in results:
        return section_name, results[section_name]
    
    # Try case-insensitive match
    for key in results.keys():
        if normalize_section_name(key) == normalized_target:
            return key, results[key]
    
    return None, None


def has_meaningful_data(df):
    """
    Check if DataFrame has meaningful data
    Returns False if:
    - DataFrame is None or empty
    - Contains only 'no rows selected' messages  
    - All values are NaN or empty strings
    """
    if df is None or df.empty:
        return False
    
    # Check for "no rows selected" messages in any column
    for col in df.columns:
        if any("no rows selected" in str(x).lower() for x in df[col]):
            return False
    
    # Check if all rows are NaN
    if df.isna().all().all():
        return False
    
    # Check if DataFrame has at least one non-empty value
    try:
        if df.astype(str).apply(lambda x: x.str.strip()).eq('').all().all():
            return False
    except:
        pass  # If conversion fails, assume there's data
    
    return True



def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), color_hex)
    
    tcPr = cell._element.get_or_add_tcPr()
    existing_shading = tcPr.find(qn('w:shd'))
    if existing_shading is not None:
        tcPr.remove(existing_shading)
        
    tcPr.append(shading_elm)

# Continue with generate_word_report and rest of your functions...


def generate_word_report(results, db_info, summary_inputs, frag_thresh, ts_thresh, custom_entries, section_order, primary_color=None):
    doc = Document()
    
    doc.add_heading("DB Health Check Report", level=0)
    # ... rest of function
    # ... rest of the function

    # Summary Report
    doc.add_heading("Summary Report", level=1)
    summary_points_order = [
        "• Locking And Blocking events:",
        "• OS Performance:",
        "• IO Performance:",
        "• Network Performance:",
        "• DB Performance parameters:",
        "• DB Maintenance:",
        "• STO (snapshot too old):",
        "• Overall Database Performance:"
    ]

    for point in summary_points_order:
        key_for_lookup = point.strip("• ").strip(":")
        text = summary_inputs.get(key_for_lookup.strip(), "N/A or Missing Input")
        p = doc.add_paragraph()
        run = p.add_run(point)
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(text)

    # Database Info
    if db_info:
        doc.add_heading("Database Information", level=1)
        for key, value in db_info.items():
            p = doc.add_paragraph()
            run_key = p.add_run(f"{key}: ")
            run_key.bold = True
            run_key.font.size = Pt(12)
            run_val = p.add_run(str(value))
            run_val.font.size = Pt(12)

    # *** KEY FIX: Create mapping with normalized section names ***
    custom_entries_map = {}
    for entry in custom_entries:
        section_name = entry.get("section", "").strip()
        if section_name:
            custom_entries_map[section_name] = entry

    # Create mapping with normalized names
    custom_entries_map = {}
    for entry in custom_entries:
        section_name = entry.get("section", "").strip()
        if section_name:
            custom_entries_map[section_name] = entry


    
    # Detail Sections - Process in the order specified
    for section in section_order:
        section = section.strip()  # Normalize
        
        # Check if this is a custom-only section (not in results)
        # UPDATED: Use case-insensitive lookup
        found_key, df = find_section_in_results(section, results)
        
        # Check if section has metadata (observations/recommendations)
        section_metadata = get_section_metadata(section)
        
        if not found_key or df is None:
            # No data found - check if we should show metadata anyway
            if section_metadata and should_show_section_without_data(section):
                # Only show sections without data if they're allowed to
                # (TABLES DEGREE and INDEX DEGREE are excluded)
                doc.add_heading(section, level=2)
                doc.add_paragraph(f"No data found for this section.")
                # Add observations/recommendations from metadata
                if section_metadata.get("observation"):
                    p = doc.add_paragraph()
                    run = p.add_run("Observation: ")
                    run.bold = True
                    run.font.size = Pt(12)
                    doc.add_paragraph(section_metadata["observation"])
                if section_metadata.get("recommendation"):
                    p = doc.add_paragraph()
                    run = p.add_run("Recommendation: ")
                    run.bold = True
                    run.font.size = Pt(12)
                    doc.add_paragraph(section_metadata["recommendation"])
                if section_metadata.get("note"):
                    p = doc.add_paragraph()
                    run = p.add_run("Note: ")
                    run.bold = True
                    run.font.size = Pt(12)
                    doc.add_paragraph(section_metadata["note"])
                continue
            # Check custom entries
            elif section in custom_entries_map:
                doc.add_heading(section, level=2)
                add_custom_entry_to_word(doc, custom_entries_map[section])
            continue
        
        # Skip sections without meaningful data (unless they have metadata)
        if not has_meaningful_data(df) and not section_metadata:
            continue
        
        # Use the found section name and dataframe
        display_df = df.copy()
        
        # Original check kept for custom entries
        if False:  # Placeholder
            if section in custom_entries_map:
                doc.add_heading(section, level=2)
                add_custom_entry_to_word(doc, custom_entries_map[section])
            continue
        
        # This section exists in results
        df = results[section]
        display_df = df.copy()

        # [Keep all the filtering logic exactly as before - lines 330-400]
        if section in ["TABLES STATS", "INDEX STATS"]:
            full_df = df.copy()
            possible_cols = ["LAST_ANALYZE_DAYS", "LAST_ANALYZED_DAYS", "DAYS_SINCE_ANALYZE"]
            col_to_use = next((c for c in possible_cols if c in full_df.columns), None)
            if col_to_use:
                full_df[col_to_use] = full_df[col_to_use].astype(str).str.extract(r"(\d+)").astype(float)
            total_count = len(full_df)
            if col_to_use:
                display_df = full_df.sort_values(by=col_to_use, ascending=False, na_position="last").head(50)
            else:
                display_df = full_df.head(50)
        elif section == "TABLE SIZE AND PARTATION":
            total_count = len(display_df)
            display_df = display_df.head(20)
        elif section == "TABLE FRAGMENTATION":
            total_count = len(display_df)
            if "WASTAGE_PERCENT" in display_df.columns:
                display_df = display_df[display_df["WASTAGE_PERCENT"] > frag_thresh]
        elif section == "INDEX FRAGMENTATION":
            total_count = len(display_df)
            if "PERCENTAGE" in display_df.columns:
                display_df = display_df[display_df["PERCENTAGE"] > frag_thresh]
        elif section == "Oracle Auto-Jobs":
            total_count = len(display_df)
            if "STATUS" in display_df.columns and not display_df.empty:
                display_df = display_df[display_df["STATUS"].str.strip().str.upper() == "ENABLED"]
                display_df = display_df.drop_duplicates(subset=["CLIENT_NAME"])
            else:
                display_df = pd.DataFrame()
        else:
            total_count = len(display_df)

        if section in ["TABLE FRAGMENTATION", "INDEX FRAGMENTATION", "TABLESPACE INFORMATION"]:
            col_check = None
            if section == "TABLE FRAGMENTATION" and "WASTAGE_PERCENT" in display_df.columns:
                col_check = "WASTAGE_PERCENT"
            elif section == "INDEX FRAGMENTATION" and "PERCENTAGE" in display_df.columns:
                col_check = "PERCENTAGE"
            elif section == "TABLESPACE INFORMATION" and "USED_PERCENT" in display_df.columns:
                col_check = "USED_PERCENT"
            if col_check:
                display_df = display_df.dropna(subset=[col_check])

        if display_df.empty:
            continue

        doc.add_heading(f"{section} (Count: {total_count})", level=2)

        if "Message" in display_df.columns and any("no rows selected" in str(x).lower() for x in display_df["Message"]):
            doc.add_paragraph("No rows selected", style="Intense Quote")
        else:
            # [Keep all table generation code exactly as before - lines 410-450]
            table = doc.add_table(rows=1, cols=len(display_df.columns))
            table.style = "Table Grid"
            for i, col in enumerate(display_df.columns):
                table.rows[0].cells[i].paragraphs[0].add_run(col).bold = True

            for _, row in display_df.iterrows():
                cells = table.add_row().cells
                is_critical_row = False
                for i, col_name in enumerate(display_df.columns):
                    cell_text = str(row[col_name])
                    run = cells[i].paragraphs[0].add_run(cell_text)
                    should_be_red_font = False
                    if section == "TABLE FRAGMENTATION" and col_name == "WASTAGE_PERCENT":
                        if row[col_name] > frag_thresh:
                            should_be_red_font = True
                            is_critical_row = True
                    elif section == "INDEX FRAGMENTATION" and col_name == "PERCENTAGE":
                        if row[col_name] > frag_thresh:
                            should_be_red_font = True
                            is_critical_row = True
                    elif section == "INVALID OBJECTS" and col_name.upper() == "STATUS":
                        if str(row[col_name]).strip().upper() == "INVALID":
                            should_be_red_font = True
                            is_critical_row = True
                    elif section == "TABLESPACE INFORMATION" and col_name.upper() == "USED_PERCENT":
                        if "USED_PERCENT" in display_df.columns and row[col_name] > ts_thresh:
                            should_be_red_font = True
                            is_critical_row = True
                    if should_be_red_font:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                excluded_sections = ["INVALID OBJECTS", "TABLE FRAGMENTATION", "INDEX FRAGMENTATION"]
                if is_critical_row and section not in excluded_sections:
                    for cell in cells:
                        set_cell_background(cell, "FFD7D7")

        # *** KEY FIX: Check if section has custom entry ***
        if section in custom_entries_map:
            # Use custom observations/recommendations
            add_custom_entry_to_word(doc, custom_entries_map[section])
        else:
            # Use default observations
            add_all_section_observations(doc, section, df, results, frag_thresh, ts_thresh, display_df)
    
    # [Keep footer code exactly as before - lines 460-520]
    doc.add_page_break()
    new_section = doc.add_section()
    footer = new_section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear()
    footer_para_line = footer.add_paragraph()
    footer_para_line.alignment = 1
    run_line = footer_para_line.add_run("_" * 80)
    run_line.font.color.rgb = RGBColor(0xe5, 0xe7, 0xeb)
    run_line.font.size = Pt(8)
    footer.add_paragraph()
    footer_para = footer.add_paragraph()
    footer_para.alignment = 1
    run1 = footer_para.add_run("Created By")
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(0x66, 0x7e, 0xea)
    run1.bold = True
    footer_para.add_run("\n")
    run2 = footer_para.add_run("Clover Infotech")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    run2.bold = True
    footer_para.add_run("\n")
    footer_para_line2 = footer.add_paragraph()
    footer_para_line2.alignment = 1
    run_line2 = footer_para_line2.add_run("_" * 80)
    run_line2.font.color.rgb = RGBColor(0xe5, 0xe7, 0xeb)
    run_line2.font.size = Pt(8)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def add_all_section_observations(doc, section, df, results, frag_thresh, ts_thresh, display_df):
    """Add ALL observations and recommendations - includes predefined metadata"""
    
    # First check if section has predefined metadata
    metadata = get_section_metadata(section)
    if metadata:
        if metadata.get("observation"):
            p = doc.add_paragraph()
            run = p.add_run("Observation: ")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(metadata["observation"])
        
        if metadata.get("recommendation"):
            p = doc.add_paragraph()
            run = p.add_run("Recommendation: ")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(metadata["recommendation"])
        
        if metadata.get("note"):
            p = doc.add_paragraph()
            run = p.add_run("Note: ")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(metadata["note"])
        
        if metadata.get("command"):
            p = doc.add_paragraph()
            run = p.add_run("Command: ")
            run.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(metadata["command"], style="Intense Quote")
        
        return  # Return after adding metadata
    
    # Below: Original hardcoded observations (kept for backward compatibility)
    
    if section == "INDEX STATS":
        total_tables = len(results.get("TABLES STATS", pd.DataFrame()))
        total_indexes = len(results.get("INDEX STATS", pd.DataFrame()))
        doc.add_paragraph(
            f"Observation: During our analysis, we observed that around **{total_tables:,}** tables and around **{total_indexes:,}** indexes "
            f"in the database have stale statistics)."
        )
        doc.add_paragraph(
            "Recommendation: We recommend performing a weekly statistics gathering during off-business hours "
            "(preferably on weekends – Saturday or Sunday) for the main application schemas, using the AUTO sampling "
            "rate to maintain balanced performance and accuracy. This proactive approach will help prevent execution plan "
            "changes and reduce the risk of suboptimal plans being generated due to stale statistics."
        )

    elif section == "TABLE FRAGMENTATION" or section == "INDEX FRAGMENTATION":
        doc.add_paragraph(
            f"Note: Defragmentation should be considered when fragmentation exceeds **{frag_thresh}%** of the object size. "
            "Excessive fragmentation leads to wasted storage, higher I/O, and slower query performance due to more blocks being scanned."
        )

    elif section == "INVALID OBJECTS":
        if "OWNER" in df.columns and "STATUS" in df.columns and not df.empty:
            invalid_df = df[df["STATUS"].str.strip().str.upper() == "INVALID"]
            invalid_count = len(invalid_df)
            if invalid_count > 0:
                owner_counts = invalid_df["OWNER"].value_counts()
                top_owner = owner_counts.index[0]
                count_top_owner = owner_counts.iloc[0]
                doc.add_paragraph(
                    f"Observation: We have found there are **{invalid_count}** invalid objects. The majority ({count_top_owner}) belong to the '{top_owner}' schema."
                )
            else:
                doc.add_paragraph("Observation: No invalid objects were found.")
        else:
            doc.add_paragraph(f"Observation: We have found there are **{len(df)}** invalid objects.")
        doc.add_paragraph("Recommendation: Recompile the invalid objects and monitor regularly to prevent application errors.")

    elif section == "TABLESPACE INFORMATION":
        doc.add_paragraph(
            f"Note: Monitor the size of tablespace which should not cross its threshold value i.e. **{ts_thresh}%**. Tablespaces over this threshold are highlighted in red."
        )

    elif section == "AUD$ TABLES":
        if not df.empty:
            if "SIZE_GB" in df.columns:
                df["SIZE_GB"] = pd.to_numeric(df["SIZE_GB"], errors="coerce")
                aud_table_size_gb = df["SIZE_GB"].sum()
            elif "SIZE_MB" in df.columns:
                df["SIZE_MB"] = pd.to_numeric(df["SIZE_MB"], errors="coerce")
                aud_table_size_gb = df["SIZE_MB"].sum() / 1024
            else:
                aud_table_size_gb = 0.0
        else:
            aud_table_size_gb = 0.0
        doc.add_paragraph(f"Observation: AUD$ table space usage is {aud_table_size_gb:.2f} GB.")
        doc.add_paragraph(
            "Recommendation: If the AUD$ table becomes large, move it to a dedicated tablespace and purge old audit records to control size and maintain performance."
        )

    elif section == "TEMP TABLESPACE STATUS":
        if not df.empty:
            col_map = {"MB_TOTAL": "TOTAL_MB", "MB_USED": "USED_MB", "MB_FREE": "FREE_MB"}
            df.rename(columns={c: col_map[c] for c in df.columns if c in col_map}, inplace=True)

            for col in ["USED_MB", "FREE_MB", "TOTAL_MB"]:
                if col in df.columns:
                    df[col] = (
                        df[col].astype(str)
                        .str.replace(",", "", regex=False)
                        .str.replace(" ", "", regex=False)
                        .str.replace("MB", "", case=False, regex=False)
                        .str.replace("GB", "", case=False, regex=False)
                        .str.strip()
                    )
                    df[col] = pd.to_numeric(df[col], errors="coerce")

            total_size_mb = df["TOTAL_MB"].sum(skipna=True) if "TOTAL_MB" in df.columns else 0
            total_used_mb = df["USED_MB"].sum(skipna=True) if "USED_MB" in df.columns else 0
            total_free_mb = df["FREE_MB"].sum(skipna=True) if "FREE_MB" in df.columns else 0

            usage_percent = (total_used_mb / total_size_mb * 100) if total_size_mb > 0 else 0

            if usage_percent < 50:
                note_text = (
                    f"Note: The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with only {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "The usage is minimal, indicating sufficient free space and no immediate concern regarding temporary space utilization."
                )
            elif usage_percent < 85:
                note_text = (
                    f"Note: The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "The usage is moderate. Regularly monitor TEMP utilization to avoid space shortages during heavy sorting operations."
                )
            else:
                note_text = (
                    f"Note: The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "TEMP usage is critically high, indicating insufficient temporary space. Immediate action is required to add TEMP files or tune large sort operations."
                )

            doc.add_paragraph(note_text)
        else:
            doc.add_paragraph("Note: TEMP tablespace details not available or report section is empty.")

    elif section == "TABLES DEGREE":
        if not df.empty:
            total_tables_parallel = len(df)
            doc.add_paragraph(
                f"Observation: We found {total_tables_parallel} tables where the degree of parallelism "
                "was set to a value other than the default (1). Higher degrees can cause Oracle to spawn "
                "multiple parallel slave processes during DML or query execution, which can increase CPU utilization "
                "and cause resource contention."
            )
            doc.add_paragraph(
                "Recommendation: Keep the degree of all tables set to 1 (default value) unless parallelism is "
                "explicitly required for large data load or analytical workloads."
            )
        else:
            doc.add_paragraph("Observation: No tables were found with degree of parallelism greater than 1.")

    elif section == "INDEX DEGREE":
        if not df.empty:
            total_indexes_parallel = len(df)
            doc.add_paragraph(
                f"Observation: We found {total_indexes_parallel} indexes where the degree of parallelism "
                "was set to a value other than the default (1). When the degree is increased, Oracle may spawn "
                "multiple parallel slave processes during query execution, which can lead to sudden session spikes "
                "and additional resource consumption on the database."
            )
            doc.add_paragraph(
                "Recommendation: It is advisable to keep the degree of all indexes set to 1 (default value), "
                "unless parallelism is explicitly required for specific workloads."
            )
        else:
            doc.add_paragraph("Observation: No indexes were found with degree of parallelism greater than 1.")

    elif section == "DATABASE DICTIONARY STATS":
        doc.add_paragraph(
            "We have found that Dictionary tables on the database don't have the latest statistics. "
            "We need to check if we can gather the Dictionary table's statistics to help internal queries. "
            "This activity should be done when there is Low load on the server."
        )
        p = doc.add_paragraph()
        run = p.add_run("Command:")
        run.bold = True
        doc.add_paragraph("EXEC DBMS_STATS.GATHER_DICTIONARY_STATS;", style="Intense Quote")
        p = doc.add_paragraph()
        run = p.add_run("Note: ")
        run.bold = True
        run_text = p.add_run("Stats gathering activity for fixed and dictionary objects should be done on quarterly basis.")
        run_text.font.size = Pt(11)

    elif section == "FIXED OBJECT STATS":
        doc.add_paragraph(
            "We have found that fixed tables on the database are not gathered. "
            "We need to check if we can gather the fixed table's statistics to help internal queries. "
            "This activity should be done when there is low load on the server."
        )
        p = doc.add_paragraph()
        run = p.add_run("Command:")
        run.bold = True
        doc.add_paragraph("EXEC DBMS_STATS.GATHER_FIXED_OBJECTS_STATS;", style="Intense Quote")
        p = doc.add_paragraph()
        run = p.add_run("Note: ")
        run.bold = True
        run_text = p.add_run("Stats gathering activity for fixed and dictionary objects should be done on quarterly basis.")
        run_text.font.size = Pt(11)

    elif section == "Oracle Auto-Jobs" and not display_df.empty:
        doc.add_paragraph(
            "Note: Oracle internal jobs are found which are running daily in background along with business jobs/sessions which is to be disabled. Need to disable below jobs on immediate basis."
        )


def add_custom_entry_to_word(doc, entry):
    """Add custom entry to Word document with Snapshot header"""
    if entry["observation"]:
        p = doc.add_paragraph()
        run = p.add_run("Observation: ")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(entry["observation"])
    
    if entry["recommendation"]:
        p = doc.add_paragraph()
        run = p.add_run("Recommendation: ")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(entry["recommendation"])
    
    if entry["note"]:
        p = doc.add_paragraph()
        run = p.add_run("Note: ")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(entry["note"])
    
    if entry["command"]:
        p = doc.add_paragraph()
        run = p.add_run("Command:")
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(entry["command"], style="Intense Quote")
    
    # Add Snapshot header and images
    if entry.get("images"):
        # Add "Snapshot" header before images
        p = doc.add_paragraph()
        run = p.add_run("Snapshot:")
        run.bold = True
        run.font.size = Pt(12)
        
        # Add images
        for img_data in entry["images"]:
            try:
                image_stream = BytesIO(img_data)
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(image_stream, width=Inches(5.0))
                para.alignment = 0
            except Exception as e:
                doc.add_paragraph(f"[Image could not be inserted: {str(e)}]")
    
    if entry.get("table"):
        try:
            table_info = entry["table"]
            headers = table_info["headers"]
            rows_data = table_info["rows"]
            
            custom_table = doc.add_table(rows=1, cols=len(headers))
            custom_table.style = "Table Grid"
            
            for idx, header in enumerate(headers):
                cell = custom_table.rows[0].cells[idx]
                cell.paragraphs[0].add_run(header).bold = True
            
            for row_data in rows_data:
                cells = custom_table.add_row().cells
                for idx, cell_value in enumerate(row_data):
                    cells[idx].paragraphs[0].add_run(str(cell_value))
        except Exception as e:
            doc.add_paragraph(f"[Table could not be inserted: {str(e)}]")


def adjust_color_brightness(hex_color, percent):
    """Adjust hex color brightness"""
    hex_color = hex_color.lstrip('#')
    rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    new_rgb = tuple(max(0, min(255, int(c * (100 + percent) / 100))) for c in rgb)
    return '#{:02x}{:02x}{:02x}'.format(*new_rgb)


# ---------------- ENHANCED HTML REPORT GENERATOR ----------------
def generate_enhanced_html_report(results, db_info, summary_inputs, frag_thresh, ts_thresh, custom_entries, section_order, primary_color="#2563eb"):
    """Generate professional-grade HTML report with modern design"""
    
    secondary_color = adjust_color_brightness(primary_color, -15)
    accent_color = adjust_color_brightness(primary_color, 30)
    
    
    
    html_parts = []
    
    # Sophisticated HTML Header with modern styling - ONLY ONCE
    html_parts.append(f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Database Health Check Report - Executive Summary</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        :root {{
            --primary: {primary_color};
            --primary-dark: {secondary_color};
            --primary-light: {accent_color};
            --success: #10b981;
            --warning: #f59e0b;
            --danger: #ef4444;
            --text-primary: #1f2937;
            --text-secondary: #6b7280;
            --bg-primary: #ffffff;
            --bg-secondary: #f9fafb;
            --border: #e5e7eb;
        }}
        
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            line-height: 1.6;
            color: var(--text-primary);
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            padding: 40px 20px;
            min-height: 100vh;
        }}
        
        .report-container {{
            max-width: 1200px;
            margin: 0 auto;
            background: var(--bg-primary);
            border-radius: 24px;
            box-shadow: 0 25px 50px -12px rgba(0, 0, 0, 0.25);
            overflow: hidden;
        }}
        
        .report-header {{
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
            color: white;
            padding: 60px 80px;
            position: relative;
            overflow: hidden;
        }}
        
        .report-header::before {{
            content: '';
            position: absolute;
            top: -50%;
            right: -50%;
            width: 200%;
            height: 200%;
            background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%);
            animation: pulse 15s ease-in-out infinite;
        }}
        
        @keyframes pulse {{
            0%, 100% {{ transform: scale(1) rotate(0deg); }}
            50% {{ transform: scale(1.1) rotate(180deg); }}
        }}
        
        .header-content {{
            position: relative;
            z-index: 1;
        }}
        
        .report-title {{
            font-size: 3em;
            font-weight: 700;
            text-align: center;
            margin-bottom: 10px;
            letter-spacing: -0.02em;
            text-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        
        .report-subtitle {{
            text-align: center;
            font-size: 1.1em;
            opacity: 0.95;
            font-weight: 400;
        }}
        
        .report-body {{
            padding: 60px 80px;
        }}
        
        .section {{
            margin-bottom: 60px;
        }}
        
        .section-header {{
            display: flex;
            align-items: center;
            margin-bottom: 30px;
            padding-bottom: 15px;
            border-bottom: 3px solid var(--primary);
        }}
        
        .section-icon {{
            width: 40px;
            height: 40px;
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            margin-right: 15px;
            color: white;
            font-weight: 700;
            font-size: 1.2em;
        }}
        
        .section-title {{
            font-size: 2em;
            font-weight: 700;
            color: var(--primary-dark);
            flex: 1;
        }}
        
        .section-count {{
            background: var(--primary-light);
            color: var(--primary-dark);
            padding: 8px 20px;
            border-radius: 50px;
            font-weight: 600;
            font-size: 0.9em;
        }}
        
        /* Executive Summary Cards */
        .summary-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 25px;
            margin-bottom: 40px;
        }}
        
        .summary-card {{
            background: linear-gradient(135deg, var(--bg-secondary) 0%, white 100%);
            border: 2px solid var(--border);
            border-radius: 16px;
            padding: 25px;
            transition: all 0.3s ease;
            position: relative;
            overflow: hidden;
        }}
        
        .summary-card::before {{
            content: '';
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 4px;
            background: linear-gradient(90deg, var(--primary) 0%, var(--primary-dark) 100%);
        }}
        
        .summary-card:hover {{
            transform: translateY(-5px);
            box-shadow: 0 20px 40px -15px rgba(0,0,0,0.2);
            border-color: var(--primary);
        }}
        
        .summary-label {{
            font-size: 0.85em;
            font-weight: 600;
            color: var(--text-secondary);
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 8px;
        }}
        
        .summary-content {{
            font-size: 1em;
            color: var(--text-primary);
            line-height: 1.5;
        }}
        
        /* Database Info Grid */
        .db-info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }}
        
        .db-info-card {{
            background: white;
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 20px;
            text-align: center;
            transition: all 0.3s ease;
        }}
        
        .db-info-card:hover {{
            border-color: var(--primary);
            box-shadow: 0 8px 16px -8px rgba(0,0,0,0.1);
        }}
        
        .db-info-label {{
            font-size: 0.8em;
            font-weight: 600;
            color: var(--text-secondary);
            text-transform: uppercase;
            letter-spacing: 0.05em;
            margin-bottom: 8px;
        }}
        
        .db-info-value {{
            font-size: 1.4em;
            font-weight: 700;
            color: var(--primary-dark);
        }}
        
        /* Modern Tables */
        .table-container {{
            overflow-x: auto;
            margin: 30px 0;
            border-radius: 12px;
            box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1);
        }}
        
        table {{
            width: 100%;
            border-collapse: separate;
            border-spacing: 0;
            background: white;
            font-size: 0.9em;
        }}
        
        thead {{
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
            color: white;
        }}
        
        th {{
            padding: 16px 12px;
            text-align: left;
            font-weight: 600;
            text-transform: uppercase;
            font-size: 0.75em;
            letter-spacing: 0.05em;
            position: sticky;
            top: 0;
            z-index: 10;
        }}
        
        th:first-child {{
            border-top-left-radius: 12px;
        }}
        
        th:last-child {{
            border-top-right-radius: 12px;
        }}
        
        td {{
            padding: 14px 12px;
            border-bottom: 1px solid var(--border);
        }}
        
        tbody tr {{
            transition: background-color 0.2s ease;
        }}
        
        tbody tr:hover {{
            background-color: var(--bg-secondary);
        }}
        
        tbody tr:last-child td:first-child {{
            border-bottom-left-radius: 12px;
        }}
        
        tbody tr:last-child td:last-child {{
            border-bottom-right-radius: 12px;
        }}
        
        .critical-row {{
            background-color: #fee2e2 !important;
        }}
        
        .critical-row:hover {{
            background-color: #fecaca !important;
        }}
        
        .red-text {{
            color: var(--danger);
            font-weight: 600;
        }}
        
        /* Insight Boxes */
        .insight-box {{
            margin: 25px 0;
            padding: 25px;
            border-radius: 12px;
            border-left: 5px solid;
            position: relative;
            overflow: hidden;
        }}
        
        .insight-box::before {{
            content: '';
            position: absolute;
            top: 0;
            right: 0;
            width: 100px;
            height: 100%;
            background: linear-gradient(90deg, transparent 0%, rgba(255,255,255,0.5) 100%);
        }}
        
        .insight-observation {{
            background: linear-gradient(135deg, #dbeafe 0%, #e0e7ff 100%);
            border-left-color: #3b82f6;
        }}
        
        .insight-recommendation {{
            background: linear-gradient(135deg, #fef3c7 0%, #fed7aa 100%);
            border-left-color: #f59e0b;
        }}
        
        .insight-note {{
            background: linear-gradient(135deg, #f3f4f6 0%, #e5e7eb 100%);
            border-left-color: #6b7280;
        }}
        
        .insight-header {{
            display: flex;
            align-items: center;
            font-weight: 700;
            font-size: 1.1em;
            margin-bottom: 12px;
        }}
        
        .insight-icon {{
            width: 24px;
            height: 24px;
            margin-right: 10px;
            display: inline-block;
        }}
        
        .insight-content {{
            line-height: 1.7;
            color: var(--text-primary);
        }}
        
        code {{
            background: #1f2937;
            color: #10b981;
            padding: 15px 20px;
            display: block;
            border-radius: 8px;
            font-family: 'Courier New', monospace;
            font-size: 0.9em;
            margin: 15px 0;
            overflow-x: auto;
            border-left: 4px solid #10b981;
        }}
        
        .no-data {{
            text-align: center;
            padding: 40px;
            color: var(--text-secondary);
            font-style: italic;
            background: var(--bg-secondary);
            border-radius: 12px;
            border: 2px dashed var(--border);
        }}
        
        /* Footer */
        .report-footer {{
            background: var(--bg-secondary);
            padding: 40px 80px;
            text-align: center;
            border-top: 1px solid var(--border);
        }}
        
        .footer-text {{
            color: var(--text-secondary);
            font-size: 0.9em;
        }}
        
        /* Print Styles */
        @media print {{
            body {{
                background: white;
                padding: 0;
            }}
            
            .report-container {{
                box-shadow: none;
                border-radius: 0;
            }}
            
            .report-header {{
                page-break-after: avoid;
            }}
            
            .section {{
                page-break-inside: avoid;
            }}
            
            table {{
                page-break-inside: auto;
            }}
            
            tr {{
                page-break-inside: avoid;
                page-break-after: auto;
            }}
        }}
        
        /* Responsive Design */
        @media (max-width: 768px) {{
            .report-header, .report-body, .report-footer {{
                padding: 30px 20px;
            }}
            
            .report-title {{
                font-size: 2em;
            }}
            
            .section-title {{
                font-size: 1.5em;
            }}
            
            .summary-grid, .db-info-grid {{
                grid-template-columns: 1fr;
            }}
        }}
    </style>
</head>
<body>
    <div class="report-container">
        <div class="report-header">
            <div class="header-content">
                <h1 class="report-title">Database Health Check Report</h1>
                <p class="report-subtitle">Executive Summary & Technical Analysis · Generated {pd.Timestamp.now().strftime('%B %d, %Y at %I:%M %p')}</p>
            </div>
        </div>
        
        <div class="report-body">
""")
    
    # Executive Summary Section
    html_parts.append("""
            <div class="section">
                <div class="section-header">
                    <div class="section-icon">📋</div>
                    <h2 class="section-title">Executive Summary</h2>
                </div>
                <div class="summary-grid">
""")
    
    # ... rest of your code continues here
    
    summary_points_order = [
        ("🔒", "Locking And Blocking events"),
        ("💻", "OS Performance"),
        ("💾", "IO Performance"),
        ("🌐", "Network Performance"),
        ("⚙️", "DB Performance parameters"),
        ("🔧", "DB Maintenance"),
        ("📸", "STO (snapshot too old)"),
        ("📊", "Overall Database Performance")
    ]
    
    for icon, key in summary_points_order:
        text = summary_inputs.get(key, "N/A")
        html_parts.append(f"""
                    <div class="summary-card">
                        <div class="summary-label">{icon} {key}</div>
                        <div class="summary-content">{text}</div>
                    </div>
""")
    
    html_parts.append("""
                </div>
            </div>
""")
    
    # Database Information Section
    if db_info:
        html_parts.append("""
            <div class="section">
                <div class="section-header">
                    <div class="section-icon">🗄️</div>
                    <h2 class="section-title">Database Information</h2>
                </div>
                <div class="db-info-grid">
""")
        
        for key, value in db_info.items():
            html_parts.append(f"""
                    <div class="db-info-card">
                        <div class="db-info-label">{key}</div>
                        <div class="db-info-value">{value}</div>
                    </div>
""")
        
        html_parts.append("""
                </div>
            </div>
""")
    
    # *** KEY FIX: Create mapping with normalized section names - SAME AS WORD ***
    custom_entries_map = {}
    for entry in custom_entries:
        section_name = entry.get("section", "").strip()
        if section_name:
            custom_entries_map[section_name] = entry
    
    section_icons = {
        "TABLES STATS": "📊",
        "INDEX STATS": "🔍",
        "TABLE FRAGMENTATION": "🧩",
        "INDEX FRAGMENTATION": "🔧",
        "TABLESPACE INFORMATION": "💽",
        "TEMP TABLESPACE STATUS": "⏱️",
        "TABLE SIZE AND PARTATION": "📏",
        "INVALID OBJECTS": "⚠️",
        "UNUSED TABLES": "🗑️",
        "FIXED OBJECT STATS": "📌",
        "DATABASE DICTIONARY STATS": "📚",
        "AUD$ TABLES": "🔐",
        "TABLES DEGREE": "🔢",
        "INDEX DEGREE": "🔢",
        "Oracle Auto-Jobs": "⚙️"
    }
    
    # Process all sections in the specified order - SAME LOGIC AS WORD
    for section in section_order:
        section = section.strip()  # Normalize
        
        # Check if this is a custom-only section (not in results)
        if section not in results:
            if section in custom_entries_map:
                icon = "📝"
                html_parts.append(f"""
            <div class="section">
                <div class="section-header">
                    <div class="section-icon">{icon}</div>
                    <h2 class="section-title">{section}</h2>
                </div>
""")
                add_custom_entry_to_html(html_parts, custom_entries_map[section])
                html_parts.append("</div>")
            continue
        
        # This section exists in results
        df = results[section]
        display_df = df.copy()
        
        # Apply same filtering logic as Word report
        if section in ["TABLES STATS", "INDEX STATS"]:
            full_df = df.copy()
            possible_cols = ["LAST_ANALYZE_DAYS", "LAST_ANALYZED_DAYS", "DAYS_SINCE_ANALYZE"]
            col_to_use = next((c for c in possible_cols if c in full_df.columns), None)
            
            if col_to_use:
                full_df[col_to_use] = full_df[col_to_use].astype(str).str.extract(r"(\d+)").astype(float)
            
            total_count = len(full_df)
            
            if col_to_use:
                display_df = full_df.sort_values(by=col_to_use, ascending=False, na_position="last").head(50)
            else:
                display_df = full_df.head(50)
        
        elif section == "TABLE SIZE AND PARTATION":
            total_count = len(display_df)
            display_df = display_df.head(20)
        elif section == "TABLE FRAGMENTATION":
            total_count = len(display_df)
            if "WASTAGE_PERCENT" in display_df.columns:
                display_df = display_df[display_df["WASTAGE_PERCENT"] > frag_thresh]
        elif section == "INDEX FRAGMENTATION":
            total_count = len(display_df)
            if "PERCENTAGE" in display_df.columns:
                display_df = display_df[display_df["PERCENTAGE"] > frag_thresh]
        elif section == "Oracle Auto-Jobs":
            total_count = len(display_df)
            if "STATUS" in display_df.columns and not display_df.empty:
                display_df = display_df[display_df["STATUS"].str.strip().str.upper() == "ENABLED"]
                display_df = display_df.drop_duplicates(subset=["CLIENT_NAME"])
            else:
                display_df = pd.DataFrame()
        else:
            total_count = len(display_df)
        
        if section in ["TABLE FRAGMENTATION", "INDEX FRAGMENTATION", "TABLESPACE INFORMATION"]:
            col_check = None
            if section == "TABLE FRAGMENTATION" and "WASTAGE_PERCENT" in display_df.columns:
                col_check = "WASTAGE_PERCENT"
            elif section == "INDEX FRAGMENTATION" and "PERCENTAGE" in display_df.columns:
                col_check = "PERCENTAGE"
            elif section == "TABLESPACE INFORMATION" and "USED_PERCENT" in display_df.columns:
                col_check = "USED_PERCENT"
            if col_check:
                display_df = display_df.dropna(subset=[col_check])
        
        if display_df.empty:
            continue
        
        icon = section_icons.get(section, "📄")
        
        html_parts.append(f"""
            <div class="section">
                <div class="section-header">
                    <div class="section-icon">{icon}</div>
                    <h2 class="section-title">{section}</h2>
                    <span class="section-count">{total_count} Records</span>
                </div>
""")
        
        if "Message" in display_df.columns and any("no rows selected" in str(x).lower() for x in display_df["Message"]):
            html_parts.append('<div class="no-data">🔭 No data available for this section</div>')
        else:
            # Build table
            html_parts.append('<div class="table-container"><table><thead><tr>')
            for col in display_df.columns:
                html_parts.append(f"<th>{col}</th>")
            html_parts.append("</tr></thead><tbody>")
            
            for _, row in display_df.iterrows():
                is_critical_row = False
                row_cells = []
                
                for col_name in display_df.columns:
                    cell_text = str(row[col_name])
                    should_be_red = False
                    
                    if section == "TABLE FRAGMENTATION" and col_name == "WASTAGE_PERCENT":
                        if row[col_name] > frag_thresh:
                            should_be_red = True
                            is_critical_row = True
                    elif section == "INDEX FRAGMENTATION" and col_name == "PERCENTAGE":
                        if row[col_name] > frag_thresh:
                            should_be_red = True
                            is_critical_row = True
                    elif section == "INVALID OBJECTS" and col_name.upper() == "STATUS":
                        if str(row[col_name]).strip().upper() == "INVALID":
                            should_be_red = True
                            is_critical_row = True
                    elif section == "TABLESPACE INFORMATION" and col_name.upper() == "USED_PERCENT":
                        if "USED_PERCENT" in display_df.columns and row[col_name] > ts_thresh:
                            should_be_red = True
                            is_critical_row = True
                    
                    if should_be_red:
                        row_cells.append(f'<td class="red-text">{cell_text}</td>')
                    else:
                        row_cells.append(f'<td>{cell_text}</td>')
                
                excluded_sections = ["INVALID OBJECTS", "TABLE FRAGMENTATION", "INDEX FRAGMENTATION"]
                if is_critical_row and section not in excluded_sections:
                    html_parts.append('<tr class="critical-row">')
                else:
                    html_parts.append('<tr>')
                html_parts.extend(row_cells)
                html_parts.append('</tr>')
            
            html_parts.append("</tbody></table></div>")
        
        # *** KEY FIX: Check if section has custom entry - SAME AS WORD ***
        if section in custom_entries_map:
            # Use custom observations/recommendations
            add_custom_entry_to_html(html_parts, custom_entries_map[section])
        else:
            # Use default observations
            add_html_section_observations(html_parts, section, df, results, frag_thresh, ts_thresh, display_df)

        html_parts.append("</div>")
    
    
    # Footer
# Footer - STUNNING ANIMATIONS VERSION:
    html_parts.append(f"""
        </div>
        
        <div class="report-footer">
            <div style="position: relative; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 50px 80px; margin-top: 40px; box-shadow: 0 -8px 32px rgba(102, 126, 234, 0.3); overflow: hidden;">
                
                <!-- Animated gradient overlay -->
                <div style="position: absolute; width: 200%; height: 200%; top: -50%; left: -50%; background: radial-gradient(circle, rgba(255,255,255,0.1) 0%, transparent 70%); animation: rotate 20s linear infinite;"></div>
                
                <!-- Multiple floating particles -->
                <div style="position: absolute; width: 100%; height: 100%; top: 0; left: 0; overflow: hidden;">
                    <div style="position: absolute; width: 250px; height: 250px; border-radius: 50%; background: radial-gradient(circle, rgba(255,255,255,0.15), transparent); top: -125px; left: 5%; animation: float 6s ease-in-out infinite;"></div>
                    <div style="position: absolute; width: 180px; height: 180px; border-radius: 50%; background: radial-gradient(circle, rgba(255,255,255,0.12), transparent); top: 50%; right: 8%; animation: float 8s ease-in-out infinite 1s;"></div>
                    <div style="position: absolute; width: 200px; height: 200px; border-radius: 50%; background: radial-gradient(circle, rgba(255,255,255,0.1), transparent); bottom: -100px; right: 15%; animation: float 7s ease-in-out infinite 2s;"></div>
                    <div style="position: absolute; width: 150px; height: 150px; border-radius: 50%; background: radial-gradient(circle, rgba(255,255,255,0.08), transparent); top: 30%; left: 12%; animation: float 9s ease-in-out infinite 0.5s;"></div>
                </div>
                
                <!-- Glowing particles -->
                <div style="position: absolute; width: 100%; height: 100%; top: 0; left: 0;">
                    <div style="position: absolute; width: 4px; height: 4px; background: white; border-radius: 50%; top: 20%; left: 15%; animation: twinkle 3s ease-in-out infinite; box-shadow: 0 0 10px white;"></div>
                    <div style="position: absolute; width: 3px; height: 3px; background: white; border-radius: 50%; top: 60%; left: 25%; animation: twinkle 4s ease-in-out infinite 1s; box-shadow: 0 0 8px white;"></div>
                    <div style="position: absolute; width: 4px; height: 4px; background: white; border-radius: 50%; top: 40%; right: 20%; animation: twinkle 3.5s ease-in-out infinite 0.5s; box-shadow: 0 0 10px white;"></div>
                    <div style="position: absolute; width: 3px; height: 3px; background: white; border-radius: 50%; top: 70%; right: 30%; animation: twinkle 4.5s ease-in-out infinite 2s; box-shadow: 0 0 8px white;"></div>
                    <div style="position: absolute; width: 4px; height: 4px; background: white; border-radius: 50%; top: 25%; right: 40%; animation: twinkle 3.2s ease-in-out infinite 1.5s; box-shadow: 0 0 10px white;"></div>
                </div>
                
                <div style="position: relative; z-index: 2; text-align: center;">
                    
                    <!-- Animated top line with glow pulse -->
                    <div style="width: 120px; height: 4px; background: white; margin: 0 auto 25px; border-radius: 3px; box-shadow: 0 0 20px rgba(255,255,255,0.8); position: relative; overflow: hidden; animation: glow-pulse 2s ease-in-out infinite;">
                        <div style="position: absolute; width: 100%; height: 100%; background: linear-gradient(90deg, transparent, rgba(255,255,255,0.9), transparent); animation: shimmer 2s infinite;"></div>
                    </div>
                    
                    <!-- Created By with fade-in -->
                    <p style="font-size: 0.8em; color: rgba(255,255,255,0.9); text-transform: uppercase; letter-spacing: 4px; font-weight: 600; margin: 0 0 18px 0; animation: fade-in-up 1s ease-out;">
                        Created By
                    </p>
                    
                    <!-- Company Name with scale animation and multiple glows -->
                    <h1 style="font-size: 2.8em; font-weight: 900; color: white; margin: 0; letter-spacing: 3px; line-height: 1.1; font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; animation: text-glow 3s ease-in-out infinite, scale-in 1s ease-out; text-shadow: 0 0 20px rgba(255,255,255,0.5), 0 0 40px rgba(255,255,255,0.3), 0 4px 20px rgba(0,0,0,0.3);">
                        COE PT TEAM
                    </h1>
                    
                    <!-- Animated bottom line with glow pulse -->
                    <div style="width: 120px; height: 4px; background: white; margin: 25px auto 0; border-radius: 3px; box-shadow: 0 0 20px rgba(255,255,255,0.8); position: relative; overflow: hidden; animation: glow-pulse 2s ease-in-out infinite 1s;">
                        <div style="position: absolute; width: 100%; height: 100%; background: linear-gradient(90deg, transparent, rgba(255,255,255,0.9), transparent); animation: shimmer 2s infinite;"></div>
                    </div>
                    
                </div>
            </div>
        </div>
        
        <!-- Advanced CSS animations -->
        <style>
            @keyframes float {{
                0%, 100% {{ transform: translateY(0px) translateX(0px); }}
                25% {{ transform: translateY(-15px) translateX(10px); }}
                50% {{ transform: translateY(-25px) translateX(-5px); }}
                75% {{ transform: translateY(-15px) translateX(10px); }}
            }}
            
            @keyframes shimmer {{
                0% {{ transform: translateX(-100%); }}
                100% {{ transform: translateX(200%); }}
            }}
            
            @keyframes rotate {{
                0% {{ transform: rotate(0deg); }}
                100% {{ transform: rotate(360deg); }}
            }}
            
            @keyframes twinkle {{
                0%, 100% {{ opacity: 0.3; transform: scale(1); }}
                50% {{ opacity: 1; transform: scale(1.5); }}
            }}
            
            @keyframes glow-pulse {{
                0%, 100% {{ box-shadow: 0 0 15px rgba(255,255,255,0.6); }}
                50% {{ box-shadow: 0 0 30px rgba(255,255,255,1), 0 0 50px rgba(255,255,255,0.6); }}
            }}
            
            @keyframes text-glow {{
                0%, 100% {{ 
                    text-shadow: 0 0 20px rgba(255,255,255,0.5), 
                                 0 0 40px rgba(255,255,255,0.3), 
                                 0 4px 20px rgba(0,0,0,0.3); 
                }}
                50% {{ 
                    text-shadow: 0 0 30px rgba(255,255,255,0.8), 
                                 0 0 60px rgba(255,255,255,0.5), 
                                 0 0 80px rgba(255,255,255,0.3),
                                 0 4px 20px rgba(0,0,0,0.3); 
                }}
            }}
            
            @keyframes fade-in-up {{
                0% {{ opacity: 0; transform: translateY(20px); }}
                100% {{ opacity: 1; transform: translateY(0); }}
            }}
            
            @keyframes scale-in {{
                0% {{ opacity: 0; transform: scale(0.8); }}
                100% {{ opacity: 1; transform: scale(1); }}
            }}
            
            @media print {{
                .report-footer {{
                    page-break-before: avoid;
                }}
                .report-footer * {{
                    animation: none !important;
                }}
            }}
        </style>
    </div>
</body>
</html>
""")
    
    return "".join(html_parts)
    



# [KEEP REST OF YOUR HELPER FUNCTIONS]
# REPLACE the add_html_section_observations function in your Checkup.py file

def add_html_section_observations(html_parts, section, df, results, frag_thresh, ts_thresh, display_df):
    """Add ALL observations to HTML - matching Word report with styled boxes"""
    
    # First check if section has predefined metadata
    metadata = get_section_metadata(section)
    if metadata:
        if metadata.get("observation"):
            html_parts.append(f"""
        <div class="insight-box insight-observation">
            <div class="insight-header">
                <span class="insight-icon">🔍</span>
                Observation
            </div>
            <div class="insight-content">
                {metadata["observation"]}
            </div>
        </div>
        """)
        
        if metadata.get("recommendation"):
            html_parts.append(f"""
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">💡</span>
                Recommendation
            </div>
            <div class="insight-content">
                {metadata["recommendation"]}
            </div>
        </div>
        """)
        
        if metadata.get("note"):
            html_parts.append(f"""
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">
                {metadata["note"]}
            </div>
        </div>
        """)
        
        if metadata.get("command"):
            html_parts.append(f"""
        <div class="insight-box insight-command">
            <div class="insight-header">
                <span class="insight-icon">⚡</span>
                Command
            </div>
            <div class="insight-content">
                <pre style="background: #1e293b; color: #e2e8f0; padding: 12px; border-radius: 6px; overflow-x: auto;">{metadata["command"]}</pre>
            </div>
        </div>
        """)
        
        return  # Return after adding metadata
    
    # Below: Original hardcoded observations (kept for backward compatibility)
    
    if section == "INDEX STATS":
        total_tables = len(results.get("TABLES STATS", pd.DataFrame()))
        total_indexes = len(results.get("INDEX STATS", pd.DataFrame()))
        html_parts.append(f"""
        <div class="insight-box insight-observation">
            <div class="insight-header">
                <span class="insight-icon">🔍</span>
                Observation
            </div>
            <div class="insight-content">
                During our analysis, we observed that around <strong>{total_tables:,}</strong> tables 
                and around <strong>{total_indexes:,}</strong> indexes in the database have stale statistics.
            </div>
        </div>
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">💡</span>
                Recommendation
            </div>
            <div class="insight-content">
                We recommend performing a weekly statistics gathering during off-business hours 
                (preferably on weekends – Saturday or Sunday) for the main application schemas, using the AUTO sampling 
                rate to maintain balanced performance and accuracy.
            </div>
        </div>
        """)
    
    elif section == "TABLE FRAGMENTATION" or section == "INDEX FRAGMENTATION":
        html_parts.append(f"""
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">
                Defragmentation should be considered when fragmentation exceeds <strong>{frag_thresh}%</strong> 
                of the object size. Excessive fragmentation leads to wasted storage, higher I/O, and slower query performance.
            </div>
        </div>
        """)
    
    elif section == "INVALID OBJECTS":
        if "OWNER" in df.columns and "STATUS" in df.columns and not df.empty:
            invalid_df = df[df["STATUS"].str.strip().str.upper() == "INVALID"]
            invalid_count = len(invalid_df)
            if invalid_count > 0:
                owner_counts = invalid_df["OWNER"].value_counts()
                top_owner = owner_counts.index[0]
                count_top_owner = owner_counts.iloc[0]
                html_parts.append(f"""
                <div class="insight-box insight-observation">
                    <div class="insight-header">
                        <span class="insight-icon">🔍</span>
                        Observation
                    </div>
                    <div class="insight-content">
                        We have found there are <strong>{invalid_count}</strong> invalid objects. 
                        The majority ({count_top_owner}) belong to the '{top_owner}' schema.
                    </div>
                </div>
                """)
            else:
                html_parts.append("""
                <div class="insight-box insight-observation">
                    <div class="insight-header">
                        <span class="insight-icon">🔍</span>
                        Observation
                    </div>
                    <div class="insight-content">
                        No invalid objects were found.
                    </div>
                </div>
                """)
        else:
            html_parts.append(f"""
            <div class="insight-box insight-observation">
                <div class="insight-header">
                    <span class="insight-icon">🔍</span>
                    Observation
                </div>
                <div class="insight-content">
                    We have found there are <strong>{len(df)}</strong> invalid objects.
                </div>
            </div>
            """)
        html_parts.append("""
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">💡</span>
                Recommendation
            </div>
            <div class="insight-content">
                Recompile the invalid objects and monitor regularly to prevent application errors.
            </div>
        </div>
        """)
    
    elif section == "TABLESPACE INFORMATION":
        html_parts.append(f"""
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">
                Monitor the size of tablespace which should not cross its threshold value i.e. 
                <strong>{ts_thresh}%</strong>. Tablespaces over this threshold are highlighted in red.
            </div>
        </div>
        """)
    
    elif section == "AUD$ TABLES":
        if not df.empty:
            df_copy = df.copy()
            if "SIZE_GB" in df_copy.columns:
                df_copy["SIZE_GB"] = pd.to_numeric(df_copy["SIZE_GB"], errors="coerce")
                aud_table_size_gb = df_copy["SIZE_GB"].sum()
            elif "SIZE_MB" in df_copy.columns:
                df_copy["SIZE_MB"] = pd.to_numeric(df_copy["SIZE_MB"], errors="coerce")
                aud_table_size_gb = df_copy["SIZE_MB"].sum() / 1024
            else:
                aud_table_size_gb = 0.0
        else:
            aud_table_size_gb = 0.0
        html_parts.append(f"""
        <div class="insight-box insight-observation">
            <div class="insight-header">
                <span class="insight-icon">🔍</span>
                Observation
            </div>
            <div class="insight-content">
                AUD$ table space usage is {aud_table_size_gb:.2f} GB.
            </div>
        </div>
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">💡</span>
                Recommendation
            </div>
            <div class="insight-content">
                If the AUD$ table becomes large, move it to a dedicated tablespace and 
                purge old audit records to control size and maintain performance.
            </div>
        </div>
        """)
    
    elif section == "TEMP TABLESPACE STATUS":
        if not df.empty:
            df_copy = df.copy()
            col_map = {"MB_TOTAL": "TOTAL_MB", "MB_USED": "USED_MB", "MB_FREE": "FREE_MB"}
            df_copy.rename(columns={c: col_map[c] for c in df_copy.columns if c in col_map}, inplace=True)
            
            for col in ["USED_MB", "FREE_MB", "TOTAL_MB"]:
                if col in df_copy.columns:
                    df_copy[col] = (
                        df_copy[col].astype(str)
                        .str.replace(",", "", regex=False)
                        .str.replace(" ", "", regex=False)
                        .str.replace("MB", "", case=False, regex=False)
                        .str.replace("GB", "", case=False, regex=False)
                        .str.strip()
                    )
                    df_copy[col] = pd.to_numeric(df_copy[col], errors="coerce")
            
            total_size_mb = df_copy["TOTAL_MB"].sum(skipna=True) if "TOTAL_MB" in df_copy.columns else 0
            total_used_mb = df_copy["USED_MB"].sum(skipna=True) if "USED_MB" in df_copy.columns else 0
            total_free_mb = df_copy["FREE_MB"].sum(skipna=True) if "FREE_MB" in df_copy.columns else 0
            
            usage_percent = (total_used_mb / total_size_mb * 100) if total_size_mb > 0 else 0
            
            if usage_percent < 50:
                note_text = (
                    f"The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with only {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "The usage is minimal, indicating sufficient free space."
                )
            elif usage_percent < 85:
                note_text = (
                    f"The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "The usage is moderate. Monitor regularly."
                )
            else:
                note_text = (
                    f"The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "TEMP usage is critically high. Immediate action required."
                )
            
            html_parts.append(f"""
            <div class="insight-box insight-note">
                <div class="insight-header">
                    <span class="insight-icon">📝</span>
                    Note
                </div>
                <div class="insight-content">
                    {note_text}
                </div>
            </div>
            """)
        else:
            html_parts.append("""
            <div class="insight-box insight-note">
                <div class="insight-header">
                    <span class="insight-icon">📝</span>
                    Note
                </div>
                <div class="insight-content">
                    TEMP tablespace details not available.
                </div>
            </div>
            """)
    
    elif section == "TABLES DEGREE":
        if not df.empty:
            total_tables_parallel = len(df)
            html_parts.append(f"""
            <div class="insight-box insight-observation">
                <div class="insight-header">
                    <span class="insight-icon">🔍</span>
                    Observation
                </div>
                <div class="insight-content">
                    We found {total_tables_parallel} tables where the degree of parallelism 
                    was set to a value other than the default (1).
                </div>
            </div>
            <div class="insight-box insight-recommendation">
                <div class="insight-header">
                    <span class="insight-icon">💡</span>
                    Recommendation
                </div>
                <div class="insight-content">
                    Keep the degree of all tables set to 1 (default value) unless parallelism is 
                    explicitly required.
                </div>
            </div>
            """)
        else:
            html_parts.append("""
            <div class="insight-box insight-observation">
                <div class="insight-header">
                    <span class="insight-icon">🔍</span>
                    Observation
                </div>
                <div class="insight-content">
                    No tables found with degree of parallelism greater than 1.
                </div>
            </div>
            """)
    
    elif section == "INDEX DEGREE":
        if not df.empty:
            total_indexes_parallel = len(df)
            html_parts.append(f"""
            <div class="insight-box insight-observation">
                <div class="insight-header">
                    <span class="insight-icon">🔍</span>
                    Observation
                </div>
                <div class="insight-content">
                    We found {total_indexes_parallel} indexes where the degree of parallelism 
                    was set to a value other than the default (1).
                </div>
            </div>
            <div class="insight-box insight-recommendation">
                <div class="insight-header">
                    <span class="insight-icon">💡</span>
                    Recommendation
                </div>
                <div class="insight-content">
                    Keep the degree of all indexes set to 1 (default value).
                </div>
            </div>
            """)
        else:
            html_parts.append("""
            <div class="insight-box insight-observation">
                <div class="insight-header">
                    <span class="insight-icon">🔍</span>
                    Observation
                </div>
                <div class="insight-content">
                    No indexes found with degree of parallelism greater than 1.
                </div>
            </div>
            """)
    
    elif section == "DATABASE DICTIONARY STATS":
        html_parts.append("""
        <div class="insight-box insight-observation">
            <div class="insight-header">
                <span class="insight-icon">🔍</span>
                Observation
            </div>
            <div class="insight-content">
                Dictionary tables don't have the latest statistics.
            </div>
        </div>
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">⚙️</span>
                Command
            </div>
            <div class="insight-content">
                <code>EXEC DBMS_STATS.GATHER_DICTIONARY_STATS;</code>
            </div>
        </div>
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">
                Stats gathering activity for fixed and dictionary objects should be done on quarterly basis.
            </div>
        </div>
        """)
    
    elif section == "FIXED OBJECT STATS":
        html_parts.append("""
        <div class="insight-box insight-observation">
            <div class="insight-header">
                <span class="insight-icon">🔍</span>
                Observation
            </div>
            <div class="insight-content">
                Fixed tables on the database are not gathered.
            </div>
        </div>
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">⚙️</span>
                Command
            </div>
            <div class="insight-content">
                <code>EXEC DBMS_STATS.GATHER_FIXED_OBJECTS_STATS;</code>
            </div>
        </div>
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">
                Stats gathering activity for fixed and dictionary objects should be done on quarterly basis.
            </div>
        </div>
        """)
    
    elif section == "Oracle Auto-Jobs" and not display_df.empty:
        html_parts.append("""
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">
                Oracle internal jobs are found which are running daily in background.
                Need to disable below jobs on immediate basis.
            </div>
        </div>
        """)

def add_custom_entry_to_html(html_parts, entry):
    """Add custom entry to HTML with Snapshot header"""
    if entry["observation"]:
        html_parts.append(f"""
        <div class="insight-box insight-observation">
            <div class="insight-header">
                <span class="insight-icon">🔍</span>
                Observation
            </div>
            <div class="insight-content">{entry["observation"]}</div>
        </div>
        """)
    
    if entry["recommendation"]:
        html_parts.append(f"""
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">💡</span>
                Recommendation
            </div>
            <div class="insight-content">{entry["recommendation"]}</div>
        </div>
        """)
    
    if entry["note"]:
        html_parts.append(f"""
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📝</span>
                Note
            </div>
            <div class="insight-content">{entry["note"]}</div>
        </div>
        """)
    
    if entry["command"]:
        html_parts.append(f"""
        <div class="insight-box insight-recommendation">
            <div class="insight-header">
                <span class="insight-icon">⚙️</span>
                Command
            </div>
            <div class="insight-content">
                <code>{entry["command"]}</code>
            </div>
        </div>
        """)
    
    # Add Snapshot header and images
    if entry.get("images"):
        html_parts.append("""
        <div class="insight-box insight-note">
            <div class="insight-header">
                <span class="insight-icon">📸</span>
                Snapshot
            </div>
            <div class="insight-content">
        """)
        
        for img_data in entry["images"]:
            try:
                image_base64 = base64.b64encode(img_data).decode('utf-8')
                html_parts.append(f"""
                <div style="text-align: left; margin: 15px 0;">
                    <img src="data:image/png;base64,{image_base64}" 
                         style="max-width: 100%; height: auto; border: 2px solid #e5e7eb; border-radius: 8px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);" 
                         alt="Snapshot Image">
                </div>
                """)
            except:
                html_parts.append('<p style="color: red;">[Image could not be inserted]</p>')
        
        html_parts.append("""
            </div>
        </div>
        """)
    
    if entry.get("table"):
        try:
            table_info = entry["table"]
            headers = table_info["headers"]
            rows_data = table_info["rows"]
            
            html_parts.append('<div class="table-container"><table><thead><tr>')
            for header in headers:
                html_parts.append(f"<th>{header}</th>")
            html_parts.append("</tr></thead><tbody>")
            
            for row_data in rows_data:
                html_parts.append("<tr>")
                for cell_value in row_data:
                    html_parts.append(f"<td>{cell_value}</td>")
                html_parts.append("</tr>")
            
            html_parts.append("</tbody></table></div>")
        except:
            html_parts.append('<p style="color: red;">[Table could not be inserted]</p>')

# Enhanced DB Health Check Analyzer - Complete UI Code
# Place this after all your helper functions


def plot_tablespace_gauges(df, ts_thresh):
    cols = st.columns(min(len(df), 4))
    
    for i, row in df.iterrows():
        name = row['TABLESPACE_NAME']
        used_pct = row['USED_PERCENT']
        
        if pd.isna(used_pct):
            cols[i % 4].warning(f"Gauge Error: {name} (Data Missing)")
            continue
            
        color = 'green'
        if used_pct >= ts_thresh:
            color = 'red'
        elif used_pct >= (ts_thresh * 0.8):
            color = 'gold'
            
        fig = go.Figure(go.Indicator(
            mode = "gauge+number",
            value = used_pct,
            title = {'text': f"{name} Used (%)"},
            gauge = {
                'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': "darkblue"},
                'bar': {'color': color},
                'steps': [
                    {'range': [0, ts_thresh * 0.8], 'color': "lightgray"},
                    {'range': [ts_thresh * 0.8, ts_thresh], 'color': "yellow"},
                    {'range': [ts_thresh, 100], 'color': "red"}
                ],
                'threshold': {
                    'line': {'color': "black", 'width': 4},
                    'thickness': 0.75,
                    'value': ts_thresh
                }
            }
        ))
        fig.update_layout(height=250, margin=dict(l=10, r=10, t=50, b=10))
        cols[i % 4].plotly_chart(fig, use_container_width=True)


def plot_bar_chart(df, x_col, y_col, title, top_n=10, color='#3498db'):
    if x_col not in df.columns or y_col not in df.columns:
        st.warning(f"Cannot generate chart: Missing '{x_col}' or '{y_col}' column.")
        return

    plot_df = df.dropna(subset=[x_col, y_col]).sort_values(by=y_col, ascending=False).head(top_n)

    fig = px.bar(plot_df, x=x_col, y=y_col, title=title, color_discrete_sequence=[color])
    fig.update_layout(xaxis_title=x_col, yaxis_title=y_col)
    st.plotly_chart(fig, use_container_width=True)


# Enhanced DB Health Check Analyzer - Main UI Code with Section Ordering
# Add this to your existing code (keep all your helper functions as they are)
# Enhanced DB Health Check Analyzer - Streamlit UI Section
# Replace everything from "# ---------------- Streamlit UI with Global Submit Button ----------------" onwards

# ---------------- Streamlit UI with Global Submit Button ----------------
# Replace the CSS section in your Checkup.py file (around lines 540-560)
# with this updated version:

st.set_page_config(
    page_title="DB Health Check Analyzer", 
    layout="wide", 
    page_icon="📊",
    initial_sidebar_state="expanded"
)

# Enhanced Custom CSS with super smooth sidebar animations
st.markdown("""
<style>
    /* Force smooth transitions on ALL sidebar elements */
    [data-testid="stSidebar"],
    [data-testid="stSidebar"] *,
    [data-testid="stSidebar"] > div,
    [data-testid="stSidebar"] > div > div,
    section[data-testid="stSidebar"],
    section[data-testid="stSidebar"] > div {
        transition: all 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
        -webkit-transition: all 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
        -moz-transition: all 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    }
    
    /* Sidebar container smooth slide */
    [data-testid="stSidebar"] {
        transition: margin-left 0.35s cubic-bezier(0.4, 0.0, 0.2, 1), 
                    transform 0.35s cubic-bezier(0.4, 0.0, 0.2, 1),
                    width 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    }
    
    /* Inner sidebar content */
    [data-testid="stSidebar"] > div:first-child {
        transition: transform 0.35s cubic-bezier(0.4, 0.0, 0.2, 1),
                    opacity 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    }
    
    /* Sidebar navigation smooth fade */
    [data-testid="stSidebarNav"] {
        transition: all 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    }
    
    /* Arrow button smooth rotation and movement */
    [data-testid="collapsedControl"],
    button[kind="header"] {
        transition: all 0.25s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    }
    
    /* Main content area smooth shift */
    .main,
    .main .block-container,
    .stApp > header,
    [data-testid="stAppViewContainer"],
    [data-testid="stAppViewContainer"] > div {
        transition: margin-left 0.35s cubic-bezier(0.4, 0.0, 0.2, 1),
                    padding-left 0.35s cubic-bezier(0.4, 0.0, 0.2, 1) !important;
    }
    
    /* Overlay smooth fade in/out */
    [data-testid="stSidebar"]::before,
    .stApp::before {
        transition: opacity 0.3s ease-in-out !important;
    }
    
    /* Existing styles below */
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border-left: 4px solid #667eea;
    }
    .section-order-item {
        background: #f8f9fa;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
        border-left: 3px solid #667eea;
    }
</style>
""", unsafe_allow_html=True)

# Then continue with your existing custom CSS...
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border-left: 4px solid #667eea;
    }
    .section-order-item {
        background: #f8f9fa;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
        border-left: 3px solid #667eea;
    }
    
    /* Button styling to match Alert Log Analyzer */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.5rem 2rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1) !important;
    }
    .stButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15) !important;
    }
    
    /* Download button styling */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.5rem 2rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1) !important;
    }
    .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15) !important;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header"><h1>📊 Database Health Check Analyzer</h1><p>Advanced Analytics & Reporting Dashboard with Section Ordering</p></div>', unsafe_allow_html=True)

# Initialize session state
# ==================== SESSION STATE INITIALIZATION - PLACE THIS EARLY ====================
# Initialize ALL session state variables at the START
if 'form_submitted' not in st.session_state:
    st.session_state.form_submitted = False
if 'primary_color' not in st.session_state:
    st.session_state.primary_color = "#3498db"
if 'frag_thresh' not in st.session_state:
    st.session_state.frag_thresh = 50
if 'ts_thresh' not in st.session_state:
    st.session_state.ts_thresh = 80
if 'summary_inputs' not in st.session_state:
    st.session_state.summary_inputs = {
        "Locking And Blocking events": "N/A",
        "OS Performance": "Good",
        "IO Performance": "Acceptable",
        "Network Performance": "Good",
        "DB Performance parameters": "Mostly healthy, minor stale stats.",
        "DB Maintenance": "Needs fragmentation clean up.",
        "STO (snapshot too old)": "N/A",
        "Overall Database Performance": "Stable"
    }
if 'custom_entries' not in st.session_state:
    st.session_state.custom_entries = []
if 'num_custom_entries' not in st.session_state:
    st.session_state.num_custom_entries = 0
if 'table_configs' not in st.session_state:
    st.session_state.table_configs = {}
if 'individual_entries' not in st.session_state:
    st.session_state.individual_entries = {}
if 'section_selections' not in st.session_state:
    st.session_state.section_selections = {}
if 'table_selections' not in st.session_state:
    st.session_state.table_selections = {}
# 🔥 CRITICAL FIX: Initialize results and db_info
if 'results' not in st.session_state:
    st.session_state.results = {}
if 'db_info' not in st.session_state:
    st.session_state.db_info = {}
if 'realtime_execution' not in st.session_state:
    st.session_state.realtime_execution = False

# Default templates for sections
SECTION_TEMPLATES = {
    "TABLES STATS": {
        "observation": "During our analysis, we observed that tables in the database have stale statistics.",
        "recommendation": "We recommend performing a weekly statistics gathering during off-business hours (preferably on weekends – Saturday or Sunday) for the main application schemas, using the AUTO sampling rate to maintain balanced performance and accuracy.",
        "note": "",
        "command": ""
    },
    "INDEX STATS": {
        "observation": "During our analysis, we observed that indexes in the database have stale statistics.",
        "recommendation": "We recommend performing a weekly statistics gathering during off-business hours for the main application schemas.",
        "note": "",
        "command": ""
    },
    "TABLE FRAGMENTATION": {
        "observation": "Table fragmentation detected in the database.",
        "recommendation": "Defragmentation should be performed for tables exceeding the threshold.",
        "note": "Excessive fragmentation leads to wasted storage, higher I/O, and slower query performance.",
        "command": ""
    },
    "INDEX FRAGMENTATION": {
        "observation": "Index fragmentation detected in the database.",
        "recommendation": "Defragmentation should be performed for indexes exceeding the threshold.",
        "note": "Excessive fragmentation leads to wasted storage, higher I/O, and slower query performance.",
        "command": ""
    },
    "INVALID OBJECTS": {
        "observation": "Invalid objects were found in the database.",
        "recommendation": "Recompile the invalid objects and monitor regularly to prevent application errors.",
        "note": "",
        "command": "SELECT 'ALTER ' || OBJECT_TYPE || ' ' || OWNER || '.' || OBJECT_NAME || ' COMPILE;' FROM DBA_OBJECTS WHERE STATUS = 'INVALID';"
    },
    "DATABASE DICTIONARY STATS": {
        "observation": "Dictionary tables on the database don't have the latest statistics.",
        "recommendation": "Gather the Dictionary table's statistics to help internal queries. This activity should be done when there is Low load on the server.",
        "note": "Stats gathering activity for fixed and dictionary objects should be done on quarterly basis.",
        "command": "EXEC DBMS_STATS.GATHER_DICTIONARY_STATS;"
    },
    "FIXED OBJECT STATS": {
        "observation": "Fixed tables on the database are not gathered.",
        "recommendation": "Gather the fixed table's statistics to help internal queries. This activity should be done when there is low load on the server.",
        "note": "Stats gathering activity for fixed and dictionary objects should be done on quarterly basis.",
        "command": "EXEC DBMS_STATS.GATHER_FIXED_OBJECTS_STATS;"
    },
    "TABLES DEGREE": {
        "observation": "Tables were found where the degree of parallelism was set to a value other than the default (1).",
        "recommendation": "Keep the degree of all tables set to 1 (default value) unless parallelism is explicitly required for large data load or analytical workloads.",
        "note": "",
        "command": ""
    },
    "INDEX DEGREE": {
        "observation": "Indexes were found where the degree of parallelism was set to a value other than the default (1).",
        "recommendation": "Keep the degree of all indexes set to 1 (default value) unless parallelism is explicitly required for specific workloads.",
        "note": "",
        "command": ""
    },
    "AUD$ TABLES": {
        "observation": "AUD$ table space usage detected.",
        "recommendation": "If the AUD$ table becomes large, move it to a dedicated tablespace and purge old audit records to control size and maintain performance.",
        "note": "",
        "command": ""
    },
    "TABLESPACE INFORMATION": {
        "observation": "Tablespace usage monitored.",
        "recommendation": "Monitor the size of tablespace to ensure it doesn't cross threshold values.",
        "note": "Tablespaces approaching or exceeding threshold should be expanded or cleaned up.",
        "command": ""
    },
    "Oracle Auto-Jobs": {
        "observation": "Oracle internal jobs are running daily in background.",
        "recommendation": "Disable unnecessary Oracle auto-jobs to reduce background load.",
        "note": "Need to disable jobs on immediate basis if they conflict with business jobs.",
        "command": ""
    }
}

# ==================== SIDEBAR ====================
# Enhanced DB Health Check Analyzer - Streamlit UI Section with Individual Save Buttons
# Replace the Custom Entries Section in your sidebar (starting from line ~540)

# ==================== HELPER FUNCTION - PLACE THIS BEFORE SIDEBAR ====================
def get_section_defaults(section, df=None, results=None, frag_thresh=50, ts_thresh=80):
    """Generate default observation and recommendation based on section type"""
    defaults = {
        "observation": "",
        "recommendation": "",
        "note": "",
        "command": ""
    }
    
    if section == "INDEX STATS":
        if results:
            total_tables = len(results.get("TABLES STATS", pd.DataFrame()))
            total_indexes = len(results.get("INDEX STATS", pd.DataFrame()))
            defaults["observation"] = f"During our analysis, we observed that around {total_tables:,} tables and around {total_indexes:,} indexes in the database have stale statistics."
        else:
            defaults["observation"] = "During our analysis, we observed that tables and indexes in the database have stale statistics."
        defaults["recommendation"] = "We recommend performing a weekly statistics gathering during off-business hours (preferably on weekends – Saturday or Sunday) for the main application schemas, using the AUTO sampling rate to maintain balanced performance and accuracy."
    
    elif section == "TABLE FRAGMENTATION":
        defaults["note"] = f"Defragmentation should be considered when fragmentation exceeds {frag_thresh}% of the object size. Excessive fragmentation leads to wasted storage, higher I/O, and slower query performance."
    
    elif section == "INDEX FRAGMENTATION":
        defaults["note"] = f"Defragmentation should be considered when fragmentation exceeds {frag_thresh}% of the object size. Excessive fragmentation leads to wasted storage, higher I/O, and slower query performance."
    
    elif section == "INVALID OBJECTS":
        if df is not None and "OWNER" in df.columns and "STATUS" in df.columns and not df.empty:
            invalid_df = df[df["STATUS"].str.strip().str.upper() == "INVALID"]
            invalid_count = len(invalid_df)
            if invalid_count > 0:
                owner_counts = invalid_df["OWNER"].value_counts()
                top_owner = owner_counts.index[0]
                count_top_owner = owner_counts.iloc[0]
                defaults["observation"] = f"We have found there are {invalid_count} invalid objects. The majority ({count_top_owner}) belong to the '{top_owner}' schema."
            else:
                defaults["observation"] = "No invalid objects were found."
        else:
            if df is not None:
                defaults["observation"] = f"We have found there are {len(df)} invalid objects."
            else:
                defaults["observation"] = "We have found invalid objects in the database."
        defaults["recommendation"] = "Recompile the invalid objects and monitor regularly to prevent application errors."
    
    elif section == "TABLESPACE INFORMATION":
        defaults["note"] = f"Monitor the size of tablespace which should not cross its threshold value i.e. {ts_thresh}%. Tablespaces over this threshold are highlighted in red."
    
    elif section == "AUD$ TABLES":
        if df is not None and not df.empty:
            df_copy = df.copy()
            if "SIZE_GB" in df_copy.columns:
                df_copy["SIZE_GB"] = pd.to_numeric(df_copy["SIZE_GB"], errors="coerce")
                aud_table_size_gb = df_copy["SIZE_GB"].sum()
            elif "SIZE_MB" in df_copy.columns:
                df_copy["SIZE_MB"] = pd.to_numeric(df_copy["SIZE_MB"], errors="coerce")
                aud_table_size_gb = df_copy["SIZE_MB"].sum() / 1024
            else:
                aud_table_size_gb = 0.0
            defaults["observation"] = f"AUD$ table space usage is {aud_table_size_gb:.2f} GB."
        else:
            defaults["observation"] = "AUD$ table space usage detected."
        defaults["recommendation"] = "If the AUD$ table becomes large, move it to a dedicated tablespace and purge old audit records to control size and maintain performance."
    
    elif section == "TEMP TABLESPACE STATUS":
        if df is not None and not df.empty:
            df_copy = df.copy()
            col_map = {"MB_TOTAL": "TOTAL_MB", "MB_USED": "USED_MB", "MB_FREE": "FREE_MB"}
            df_copy.rename(columns={c: col_map[c] for c in df_copy.columns if c in col_map}, inplace=True)
            
            for col in ["USED_MB", "FREE_MB", "TOTAL_MB"]:
                if col in df_copy.columns:
                    df_copy[col] = (
                        df_copy[col].astype(str)
                        .str.replace(",", "", regex=False)
                        .str.replace(" ", "", regex=False)
                        .str.replace("MB", "", case=False, regex=False)
                        .str.replace("GB", "", case=False, regex=False)
                        .str.strip()
                    )
                    df_copy[col] = pd.to_numeric(df_copy[col], errors="coerce")
            
            total_size_mb = df_copy["TOTAL_MB"].sum(skipna=True) if "TOTAL_MB" in df_copy.columns else 0
            total_used_mb = df_copy["USED_MB"].sum(skipna=True) if "USED_MB" in df_copy.columns else 0
            total_free_mb = df_copy["FREE_MB"].sum(skipna=True) if "FREE_MB" in df_copy.columns else 0
            
            usage_percent = (total_used_mb / total_size_mb * 100) if total_size_mb > 0 else 0
            
            if usage_percent < 50:
                defaults["note"] = (
                    f"The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with only {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "The usage is minimal, indicating sufficient free space."
                )
            elif usage_percent < 85:
                defaults["note"] = (
                    f"The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "The usage is moderate. Monitor regularly."
                )
            else:
                defaults["note"] = (
                    f"The TEMP tablespace has a total size of {total_size_mb:,.2f} MB, "
                    f"with {total_used_mb:,.0f} MB used and {total_free_mb:,.2f} MB free. "
                    "TEMP usage is critically high. Immediate action required."
                )
        else:
            defaults["note"] = "TEMP tablespace details not available."
    
    elif section == "TABLES DEGREE":
        if df is not None and not df.empty:
            total_tables_parallel = len(df)
            defaults["observation"] = (
                f"We found {total_tables_parallel} tables where the degree of parallelism "
                "was set to a value other than the default (1). Higher degrees can cause Oracle to spawn "
                "multiple parallel slave processes which can increase CPU utilization."
            )
        else:
            defaults["observation"] = "No tables were found with degree of parallelism greater than 1."
        defaults["recommendation"] = (
            "Keep the degree of all tables set to 1 (default value) unless parallelism is "
            "explicitly required for large data load or analytical workloads."
        )
    
    elif section == "INDEX DEGREE":
        if df is not None and not df.empty:
            total_indexes_parallel = len(df)
            defaults["observation"] = (
                f"We found {total_indexes_parallel} indexes where the degree of parallelism "
                "was set to a value other than the default (1)."
            )
        else:
            defaults["observation"] = "No indexes were found with degree of parallelism greater than 1."
        defaults["recommendation"] = (
            "Keep the degree of all indexes set to 1 (default value), "
            "unless parallelism is explicitly required for specific workloads."
        )
    
    elif section == "DATABASE DICTIONARY STATS":
        defaults["observation"] = (
            "Dictionary tables on the database don't have the latest statistics. "
            "This activity should be done when there is low load on the server."
        )
        defaults["command"] = "EXEC DBMS_STATS.GATHER_DICTIONARY_STATS;"
        defaults["note"] = "Stats gathering activity for fixed and dictionary objects should be done on quarterly basis."
    
    elif section == "FIXED OBJECT STATS":
        defaults["observation"] = (
            "Fixed tables on the database are not gathered. "
            "This activity should be done when there is low load on the server."
        )
        defaults["command"] = "EXEC DBMS_STATS.GATHER_FIXED_OBJECTS_STATS;"
        defaults["note"] = "Stats gathering activity for fixed and dictionary objects should be done on quarterly basis."
    
    elif section == "Oracle Auto-Jobs":
        defaults["note"] = "Oracle internal jobs are found which are running daily in background. Need to disable these jobs on immediate basis."
    
    return defaults


# Replace the sidebar section in your Checkup.py (around line 2490-2550)
# with this fixed version:

with st.sidebar:
    st.header("⚙️ Configuration Panel")
    
    # FORM 1: Main Configuration
    with st.form(key="config_form_main", clear_on_submit=False):

        st.subheader("🎨 Branding & Customization")
        
        primary_color = st.color_picker("Primary Brand Color", st.session_state.primary_color)
        
        st.markdown("---")
        
        st.subheader("⚙️ Threshold Configuration")
        frag_thresh = st.slider("Fragmentation Threshold (%)", 10, 90, st.session_state.frag_thresh, 5)
        ts_thresh = st.slider("Tablespace Usage Threshold (%)", 50, 95, st.session_state.ts_thresh, 5)
        
        st.markdown("---")
        
        st.subheader("📝 Summary Report Inputs")
        summary_inputs = {}
        for k, v in st.session_state.summary_inputs.items():
            summary_inputs[k] = st.text_area(k, v, key=f"summary_{k.replace(' ', '_')}", height=70)
        
        submit_button = st.form_submit_button("🚀 Apply Configuration", use_container_width=True, type="primary")
        
        if submit_button:
            st.session_state.form_submitted = True
            
            # Save configuration to session state
            st.session_state.primary_color = primary_color
            st.session_state.frag_thresh = frag_thresh
            st.session_state.ts_thresh = ts_thresh
            st.session_state.summary_inputs = summary_inputs
            st.success("✅ Configuration applied successfully!")
    
    st.markdown("---")
    
    # Custom Entries Section - COMPLETELY FIXED VERSION
    st.subheader("➕ Custom Sections")
    
    # Use on_change callback to handle value changes properly
    def update_num_entries():
        new_value = st.session_state.num_entries_input
        
        if 'individual_entries' not in st.session_state:
            st.session_state.individual_entries = {}
        
        # Initialize new entries
        for i in range(new_value):
            if i not in st.session_state.table_configs:
                st.session_state.table_configs[i] = {"num_rows": 3, "num_cols": 3}
            if i not in st.session_state.individual_entries:
                st.session_state.individual_entries[i] = {}
        
        # Clean up removed entries
        keys_to_remove = [k for k in st.session_state.individual_entries.keys() if k >= new_value]
        for k in keys_to_remove:
            del st.session_state.individual_entries[k]
            if k in st.session_state.table_configs:
                del st.session_state.table_configs[k]
        
        # Update the main counter
        st.session_state.num_custom_entries = new_value
    
    # Direct number input with on_change callback
    num_custom_entries = st.number_input(
        "Number of Custom Entries", 
        min_value=0,
        max_value=10,
        value=st.session_state.num_custom_entries, 
        step=1,
        key="num_entries_input",
        on_change=update_num_entries,
        help="Set how many custom sections you want to add"
    )
    
    st.markdown("---")

# Initialize session state
if 'individual_entries' not in st.session_state:
    st.session_state.individual_entries = {}
if 'section_selections' not in st.session_state:
    st.session_state.section_selections = {}
if 'table_selections' not in st.session_state:
    st.session_state.table_selections = {}

available_sections = [
    "TABLES STATS", "INDEX STATS", "TABLE FRAGMENTATION", "INDEX FRAGMENTATION",
    "INVALID OBJECTS", "DATABASE DICTIONARY STATS", "FIXED OBJECT STATS", 
    "TABLES DEGREE", "INDEX DEGREE", "AUD$ TABLES", "TABLESPACE INFORMATION",
    "TEMP TABLESPACE STATUS", "TABLE SIZE AND PARTATION", "UNUSED TABLES",
    "Oracle Auto-Jobs", "CUSTOM SECTION"
]

# INDIVIDUAL CUSTOM ENTRIES
if st.session_state.num_custom_entries > 0:
    
    for i in range(st.session_state.num_custom_entries):
        with st.sidebar:
            st.markdown(f"### 📋 Entry #{i+1}")
            
            # Callback function for section change
            def on_section_change(entry_idx=i):
                new_section = st.session_state[f"section_select_{entry_idx}"]
                old_section = st.session_state.section_selections.get(entry_idx)
                
                # Only clear if actually changed
                if new_section != old_section:
                    st.session_state.section_selections[entry_idx] = new_section
                    # Clear saved entry when section changes to load new defaults
                    if entry_idx in st.session_state.individual_entries:
                        del st.session_state.individual_entries[entry_idx]
            
            # Section Selection OUTSIDE form for immediate response
            default_idx = 0
            if i in st.session_state.section_selections:
                prev = st.session_state.section_selections[i]
                if prev in available_sections:
                    default_idx = available_sections.index(prev)
            
            selected_section = st.selectbox(
                f"Section for Entry #{i+1}", 
                available_sections,
                index=default_idx,
                key=f"section_select_{i}",
                on_change=on_section_change,
                kwargs={"entry_idx": i}
            )
            
            # Ensure state is set
            if i not in st.session_state.section_selections:
                st.session_state.section_selections[i] = selected_section
            
            custom_section_name = selected_section
            
            # Custom Section Name OUTSIDE form
            if selected_section == "CUSTOM SECTION":
                # Get saved custom name if exists
                saved_custom_name = ""
                if i in st.session_state.individual_entries:
                    saved_section = st.session_state.individual_entries[i].get("section", "")
                    if saved_section and saved_section != "CUSTOM SECTION":
                        saved_custom_name = saved_section
                
                custom_section_name_input = st.text_input(
                    f"Custom Section Name", 
                    value=saved_custom_name,
                    placeholder="Enter custom section name...",
                    key=f"custom_name_{i}",
                    help="Enter a unique name for this custom section"
                )
                
                if custom_section_name_input and custom_section_name_input.strip() != "":
                    custom_section_name = custom_section_name_input
                else:
                    custom_section_name = f"CUSTOM SECTION {i+1}"
            
            # Table checkbox OUTSIDE form for immediate response
            create_table = st.checkbox(
                f"📊 Add Table to Entry #{i+1}", 
                value=st.session_state.table_selections.get(i, False),
                key=f"table_checkbox_{i}"
            )
            
            # Update state immediately
            st.session_state.table_selections[i] = create_table
            
            # Get defaults based on selected section
            df = None
            results_data = None
            if 'results' in st.session_state and selected_section in st.session_state.results:
                df = st.session_state.results[selected_section]
                results_data = st.session_state.results
            
            section_defaults = get_section_defaults(
                selected_section, 
                df, 
                results_data, 
                st.session_state.frag_thresh, 
                st.session_state.ts_thresh
            )
            
            # Load saved data if exists, otherwise use section defaults
            saved_entry = st.session_state.individual_entries.get(i, {})
            default_obs = saved_entry.get("observation", section_defaults.get("observation", ""))
            default_rec = saved_entry.get("recommendation", section_defaults.get("recommendation", ""))
            default_note = saved_entry.get("note", section_defaults.get("note", ""))
            default_cmd = saved_entry.get("command", section_defaults.get("command", ""))
            
            # Wrap entry fields in form to prevent refresh on typing
            with st.form(key=f"custom_entry_form_{i}", clear_on_submit=False):
                with st.expander(f"Configure Entry #{i+1}", expanded=True):
                    
                    observation = st.text_area(
                        f"Observation", 
                        value=default_obs,
                        key=f"obs_{i}", 
                        height=120,
                        help="Enter or modify the observation for this section"
                    )
                    
                    recommendation = st.text_area(
                        f"Recommendation", 
                        value=default_rec,
                        key=f"rec_{i}", 
                        height=120,
                        help="Enter or modify the recommendation for this section"
                    )
                    
                    note = st.text_area(
                        f"Note", 
                        value=default_note,
                        key=f"note_{i}", 
                        height=120,
                        help="Enter or modify additional notes"
                    )
                    
                    command = st.text_area(
                        f"Command", 
                        value=default_cmd,
                        key=f"cmd_{i}", 
                        height=100,
                        help="Enter SQL commands or scripts"
                    )
                    
                    st.markdown("---")
                    st.markdown("**📸 Snapshot**")
                    
                    uploaded_images = st.file_uploader(
                        f"Upload Images", 
                        type=["png", "jpg", "jpeg", "gif"], 
                        key=f"img_{i}",
                        accept_multiple_files=True,
                        help="Upload screenshots or diagrams"
                    )
                    
                    st.markdown("---")
                    
                    # Table configuration inside form (will show/hide based on checkbox)
                    table_data = None
                    if create_table:
                        st.markdown("**📊 Table Configuration:**")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            if i not in st.session_state.table_configs:
                                st.session_state.table_configs[i] = {"num_rows": 2, "num_cols": 2}
                            
                            num_rows = st.number_input(
                                f"Rows", 
                                min_value=1,
                                max_value=20,
                                value=st.session_state.table_configs[i]["num_rows"],
                                step=1,
                                key=f"rows_{i}"
                            )
                        
                        with col2:
                            num_cols = st.number_input(
                                f"Columns", 
                                min_value=1,
                                max_value=10,
                                value=st.session_state.table_configs[i]["num_cols"],
                                step=1,
                                key=f"cols_{i}"
                            )
                        
                        st.markdown("**Table Headers:**")
                        header_cols = st.columns(num_cols)
                        headers = []
                        for col_idx in range(num_cols):
                            with header_cols[col_idx]:
                                header = st.text_input(
                                    f"Col {col_idx+1}", 
                                    f"Column {col_idx+1}", 
                                    key=f"header_{i}_{col_idx}",
                                    label_visibility="collapsed"
                                )
                                headers.append(header)
                        
                        st.markdown("**Table Data:**")
                        rows_data = []
                        for row_idx in range(num_rows):
                            st.markdown(f"*Row {row_idx+1}:*")
                            row_cols = st.columns(num_cols)
                            row = []
                            for col_idx in range(num_cols):
                                with row_cols[col_idx]:
                                    cell_value = st.text_input(
                                        f"R{row_idx+1}C{col_idx+1}", 
                                        "", 
                                        key=f"cell_{i}_{row_idx}_{col_idx}",
                                        label_visibility="collapsed"
                                    )
                                    row.append(cell_value)
                            rows_data.append(row)
                        
                        table_data = {"headers": headers, "rows": rows_data}
                    
                    st.markdown("---")
                
                # Save Button inside form
                save_button = st.form_submit_button(
                    f"💾 Save Entry #{i+1}", 
                    use_container_width=True,
                    type="primary"
                )
                
                if save_button:
                    # Process images
                    images_data = []
                    if uploaded_images:
                        for img in uploaded_images:
                            images_data.append(img.read())
                    
                    # Update table config
                    if create_table:
                        st.session_state.table_configs[i]["num_rows"] = num_rows
                        st.session_state.table_configs[i]["num_cols"] = num_cols
                    
                    entry_data = {
                        "section": custom_section_name,
                        "observation": observation,
                        "recommendation": recommendation,
                        "note": note,
                        "command": command,
                        "images": images_data,
                        "table": table_data
                    }
                    
                    st.session_state.individual_entries[i] = entry_data
                    st.success(f"✅ Entry #{i+1} saved successfully!")
            
            # Show saved status outside form
            if i in st.session_state.individual_entries:
                st.success("✓ Saved")
            
            st.markdown("---")
    
    st.session_state.custom_entries = []
    for i in range(st.session_state.num_custom_entries):
        if i in st.session_state.individual_entries:
            entry = st.session_state.individual_entries[i]
            if any([entry.get("observation"), entry.get("recommendation"), 
                   entry.get("note"), entry.get("command"), 
                   entry.get("images"), entry.get("table")]):
                st.session_state.custom_entries.append(entry)
    
    if st.session_state.custom_entries:
        with st.sidebar:
            st.success(f"✅ Total Saved Entries: {len(st.session_state.custom_entries)}")
            
            with st.expander("📋 View Saved Entries Summary"):
                for idx, entry in enumerate(st.session_state.custom_entries):
                    st.markdown(f"**{idx+1}. {entry['section']}**")
                    if entry.get("observation"):
                        st.caption(f"✓ Observation")
                    if entry.get("recommendation"):
                        st.caption(f"✓ Recommendation")
                    if entry.get("note"):
                        st.caption(f"✓ Note")
                    if entry.get("command"):
                        st.caption(f"✓ Command")
                    if entry.get("images"):
                        st.caption(f"✓ {len(entry['images'])} Image(s)")
                    if entry.get("table"):
                        st.caption(f"✓ Table ({len(entry['table']['rows'])} rows)")
                    st.markdown("---")

else:
    with st.sidebar:
        st.info("👆 Set the number of custom entries above to start adding custom sections.")

# ==================== MAIN CONTENT AREA ====================
# (Keep the rest of your main content code exactly as it is)

# ==================== MAIN CONTENT AREA ====================
# ==================== MAIN CONTENT AREA ====================
st.markdown("---")

# Initialize session state for mode selection
if 'data_source_mode' not in st.session_state:
    st.session_state.data_source_mode = None

# UPDATED: Show data source selection ONLY if no mode is selected
if st.session_state.data_source_mode is None:
    st.header("🎯 Choose Data Source")
    
    # Create two clear options
    col1, col2 = st.columns(2)
    
    with col1:
        option_1 = st.button(
            "📄 Upload HTML Report",
            use_container_width=True,
            type="primary",
            help="Upload an existing HTML health check report for analysis"
        )
    
    with col2:
        option_2 = st.button(
            "🔌 Connect to Database",
            use_container_width=True,
            type="primary",
            help="Connect to live database and run health check script"
        )
    
    if option_1:
        st.session_state.data_source_mode = 'html_upload'
        st.rerun()
    
    if option_2:
        st.session_state.data_source_mode = 'database'
        # Clear any uploaded HTML data
        if 'results' in st.session_state:
            del st.session_state.results
        st.rerun()
    
    st.markdown("---")
    st.info("👆 Please select a data source option above to continue")

# ==================== OPTION 1: HTML UPLOAD ====================
uploaded_file = None  # Initialize

if st.session_state.data_source_mode == 'html_upload':
    # Show header with option to change mode
    col_header, col_change = st.columns([3, 1])
    with col_header:
        st.subheader("📄 Upload HTML Report")
    with col_change:
        if st.button("🔄 Change Mode", use_container_width=True):
            st.session_state.data_source_mode = None
            if 'results' in st.session_state:
                del st.session_state.results
            if 'db_info' in st.session_state:
                del st.session_state.db_info
            st.rerun()
    
    st.markdown("---")
    
    uploaded_file = st.file_uploader(
        "Upload DB Health Check Report (HTML)",
        type=["html", "htm"],
        help="Upload your existing HTML health check report"
    )
    
    if uploaded_file:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
    
    # Reset database mode flags
    if st.session_state.get('realtime_execution', False):
        st.session_state.realtime_execution = False

# ==================== OPTION 2: DATABASE CONNECTION + SCRIPT ====================
elif st.session_state.data_source_mode == 'database':
    # Show header with option to change mode
    col_header, col_change = st.columns([3, 1])
    with col_header:
        st.subheader("🔌 Connect to Database & Execute Script")
    with col_change:
        if st.button("🔄 Change Mode", use_container_width=True):
            st.session_state.data_source_mode = None
            if 'results' in st.session_state:
                del st.session_state.results
            if 'db_info' in st.session_state:
                del st.session_state.db_info
            if 'db_connected' in st.session_state:
                if st.session_state.get('db_connection'):
                    try:
                        st.session_state.db_connection.close()
                    except:
                        pass
                st.session_state.db_connected = False
                st.session_state.db_connection = None
            st.rerun()
    
    st.markdown("---")
    
    # Step 1: Upload SQL Script
    with st.expander("📂 Step 1: Upload Health Check Script", expanded=True):
        script_file = st.file_uploader(
            "Upload CHECKUP_NEW.txt or SQL Script",
            type=["txt", "sql"],
            help="Upload your health check SQL script"
        )
        
        if script_file:
            st.success(f"✅ Script loaded: {script_file.name}")
            
            # Save to session state
            if 'uploaded_script' not in st.session_state or st.session_state.uploaded_script != script_file.name:
                import tempfile
                import os
                temp_script_path = os.path.join(tempfile.gettempdir(), script_file.name)
                with open(temp_script_path, 'wb') as f:
                    f.write(script_file.read())
                st.session_state.uploaded_script = script_file.name
                st.session_state.uploaded_script_path = temp_script_path
    
    st.markdown("---")
    
    # Step 2: Database Connection
    with st.expander("🔌 Step 2: Connect to Database", expanded=True):
        is_connected, db_connection = render_database_connection_ui()
    
    st.markdown("---")
    
    # Step 3: Execute (only if both script and connection ready)
    if script_file and is_connected and db_connection:
        with st.expander("🚀 Step 3: Execute Health Check", expanded=True):
            render_health_check_execution_ui(db_connection)
            
            # Add note about scrolling down
            if st.session_state.get('realtime_execution') and st.session_state.get('results'):
                st.markdown("---")
                st.info("👇 **Scroll down** to view the complete interactive dashboard with all features!")
    
    elif not script_file:
        st.info("👆 Please upload SQL script first")
    elif not is_connected:
        st.info("👆 Please connect to database")



# ==================== ANALYSIS SECTION - FIXED ====================
# ==================== ANALYSIS SECTION - FIXED ====================
st.markdown("---")

# Initialize variables
has_data = False
results = None
db_info = None
data_source = None

# 🔥 CHECK 1: New HTML upload - PROCESS IMMEDIATELY
if uploaded_file is not None:
    try:
        # Show processing message
        with st.spinner("🔄 Analyzing uploaded HTML report..."):
            # Clear cache to ensure fresh analysis
            if hasattr(analyze_report, 'clear'):
                analyze_report.clear()
            
            # Read HTML content
            html_content = uploaded_file.read().decode("utf-8", errors="ignore")
            
            # Import analyze_report function
            from analysis_utils import analyze_report
            
            # Analyze the HTML
            results, db_info = analyze_report(html_content)
        
        # ✅ CRITICAL FIX: Check if results are valid
        if results and isinstance(results, dict) and len(results) > 0:
            # Store in session state
            st.session_state.results = results
            st.session_state.db_info = db_info if db_info else {}
            st.session_state.realtime_execution = False
            st.session_state.data_source_mode = 'html_upload'
            
            has_data = True
            data_source = "HTML Upload"
            
            st.success(f"✅ HTML Report Analyzed Successfully!")
        
        else:
            # Analysis returned empty results
            st.error("❌ HTML file uploaded but no data could be extracted")
            st.warning("⚠️ This could mean:")
            st.write("- The HTML file doesn't match the expected format")
            st.write("- The file contains no data tables")
            st.write("- The parsing logic needs adjustment")
            
            # Debug info
            with st.expander("🔍 Debug Information"):
                st.write(f"**Results type:** {type(results)}")
                st.write(f"**Number of sections:** {len(results) if results else 0}")
                
                # Show HTML preview
                st.write("**HTML Preview (first 1000 characters):**")
                st.code(html_content[:1000], language="html")
        
    except Exception as e:
        st.error(f"❌ Failed to analyze HTML file: {str(e)}")
        
        import traceback
        with st.expander("🔍 Error Details"):
            st.code(traceback.format_exc(), language="text")

# 🔥 CHECK 2: Session state results (from previous upload or DB execution)
elif 'results' in st.session_state:
    results = st.session_state.results
    
    # Verify it's valid and not empty
    if results and isinstance(results, dict) and len(results) > 0:
        db_info = st.session_state.get('db_info', {})
        has_data = True
        
        # Determine source
        if st.session_state.get('realtime_execution', False):
            data_source = "Live DB"
            st.markdown("## 🎯 Live Database Health Check Results")
            st.success("✅ **CONNECTED TO LIVE DATABASE** - Real-time analysis complete!")
        else:
            data_source = "HTML Upload"
            st.markdown("## 📊 HTML Report Analysis Results")
            st.success("✅ **HTML REPORT ANALYZED** - Interactive dashboard ready!")

# ==================== DISPLAY SECTION ====================
if has_data and results is not None and len(results) > 0:
    
    st.markdown("---")
    
    # Database Information
    if db_info:
        st.subheader("🖥️ Database Information")
    
        # Create responsive grid
        cols = st.columns(min(len(db_info), 4))
    
        for i, (key, val) in enumerate(db_info.items()):
            with cols[i % len(cols)]:
                with st.container():
                    st.markdown(f"""
                    <div style="
                        border: 1px solid #dee2e6;
                        border-left: 4px solid #667eea;
                        border-radius: 8px;
                        padding: 15px;
                        margin: 5px 0;
                        transition: all 0.3s ease;
                        cursor: default;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
                    " onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 8px rgba(0,0,0,0.1)';" 
                       onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='0 2px 4px rgba(0,0,0,0.05)';">
                        <div style="
                            font-size: 0.7em;
                            font-weight: 600;
                            color: #6c757d;
                            text-transform: uppercase;
                            letter-spacing: 0.5px;
                            margin-bottom: 8px;
                            white-space: nowrap;
                            overflow: hidden;
                            text-overflow: ellipsis;
                        " title="{key}">{key}</div>
                        <div style="
                            font-size: 1.1em;
                            font-weight: 700;
                            color: #2c3e50;
                            white-space: normal;
                            word-wrap: break-word;
                            line-height: 1.3;
                        " title="{val}">{val}</div>
                    </div>
                    """, unsafe_allow_html=True)
    
        st.markdown("---")
    
    # Section Order Configuration
    st.subheader("📋 Section Order Configuration")
    
    col_order1, col_order2 = st.columns([2, 1])
    
    with col_order1:
        st.info("💡 **Customize Report Order**: Drag and rearrange sections below")
    
    with col_order2:
        reset_order = st.button("🔄 Reset to Default Order", use_container_width=True)
    
    default_section_order = [
        "TABLES STATS", "INDEX STATS", "TABLE FRAGMENTATION", "INDEX FRAGMENTATION",
        "INVALID OBJECTS", "DATABASE DICTIONARY STATS", "FIXED OBJECT STATS", 
        "TABLES DEGREE", "INDEX DEGREE", "AUD$ TABLES", "TABLESPACE INFORMATION",
        "TEMP TABLESPACE STATUS", "TABLE SIZE AND PARTATION", "UNUSED TABLES",
        "Oracle Auto-Jobs"
    ]
    
    if 'section_order' not in st.session_state or reset_order:
        # UPDATED: Use case-insensitive matching and include sections based on data/metadata rules
        st.session_state.section_order = []
        for s in default_section_order:
            found_key, found_df = find_section_in_results(s, results)
            has_data = found_key and has_meaningful_data(found_df)
            has_metadata = get_section_metadata(s) is not None
            
            # Determine if section should be included
            should_include = False
            
            if has_data:
                # Always include if has data
                should_include = True
            elif has_metadata and should_show_section_without_data(s):
                # Include if has metadata AND is allowed to show without data
                # (TABLES DEGREE and INDEX DEGREE are excluded here)
                should_include = True
            
            if should_include:
                # Use the actual key from results if found, otherwise use default name
                section_key = found_key if found_key else s
                if section_key not in st.session_state.section_order:
                    st.session_state.section_order.append(section_key)
    
    # Get all available sections
    available_report_sections = list(results.keys())
    
    # Build complete list of ALL custom section names
    all_custom_section_names = []
    if st.session_state.custom_entries:
        for entry in st.session_state.custom_entries:
            custom_section = entry.get("section", "").strip()
            if custom_section:
                all_custom_section_names.append(custom_section)
                if custom_section not in available_report_sections:
                    available_report_sections.append(custom_section)
    
    # Add sections that aren't in the order yet (only if they're in default_section_order or custom)
    for section in available_report_sections:
        # Only add if it's in default_section_order or is a custom section
        if section not in st.session_state.section_order:
            if section in default_section_order or section in all_custom_section_names:
                st.session_state.section_order.append(section)
    
    # Clean up removed sections
    st.session_state.section_order = [s for s in st.session_state.section_order if s in available_report_sections]
    
    with st.expander("🎯 Customize Section Order", expanded=False):
        for idx, section in enumerate(st.session_state.section_order):
            col1, col2, col3, col4 = st.columns([0.5, 3, 0.5, 0.5])
            
            with col1:
                st.markdown(f"**{idx + 1}.**")
            
            with col2:
                if section in results:
                    df = results[section]
                    if df.empty:
                        status = "📘"
                    else:
                        status = "🟢"
                    st.markdown(f'{status} {section}')
                else:
                    st.markdown(f'🟡 {section} (Custom)')
            
            with col3:
                if idx > 0:
                    if st.button("⬆️", key=f"up_{section}_{idx}"):
                        st.session_state.section_order[idx], st.session_state.section_order[idx-1] = \
                            st.session_state.section_order[idx-1], st.session_state.section_order[idx]
                        st.rerun()
            
            with col4:
                if idx < len(st.session_state.section_order) - 1:
                    if st.button("⬇️", key=f"down_{section}_{idx}"):
                        st.session_state.section_order[idx], st.session_state.section_order[idx+1] = \
                            st.session_state.section_order[idx+1], st.session_state.section_order[idx]
                        st.rerun()
    
    section_order = st.session_state.section_order
    
    st.success(f"✅ Report will be generated with {len(section_order)} sections in your custom order")
    
    st.markdown("---")
    
    # Interactive Section Viewer
    st.subheader("📊 Interactive Section Viewer")
    
    selected_sections = st.multiselect(
        "🔍 Filter Sections",
        section_order,
        default=section_order,
        help="Select sections to display"
    )
    
    for section in selected_sections:
        # UPDATED: Use case-insensitive lookup
        found_key, df = find_section_in_results(section, results)
        
        if not found_key or df is None:
            continue
        
        # UPDATED: Skip sections without meaningful data
        if not has_meaningful_data(df):
            continue
        
        section_position = section_order.index(found_key) + 1
        count = len(df)
        
        with st.expander(f"**#{section_position}. {section}** ({count} Records)", expanded=False):
            st.dataframe(df.head(50), use_container_width=True, height=400)
            
            csv = df.to_csv(index=False).encode('utf-8')
            st.download_button(
                f"📥 Download {section} CSV",
                csv,
                f"{section.replace(' ', '_')}.csv",
                "text/csv",
                key=f"dl_{section}"
            )
    
    st.markdown("---")
    
    # Download Reports Section
    st.subheader("📥 Download Complete Reports")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Excel Report
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
            for section in section_order:
                if section in results:
                    # Sanitize sheet name by removing invalid Excel characters []:*?/\
                    safe_sheet_name = section.replace(':', '-').replace('[', '(').replace(']', ')').replace('*', '').replace('?', '').replace('/', '-').replace('\\', '-')
                    results[section].to_excel(writer, sheet_name=safe_sheet_name[:31], index=False)
        excel_buffer.seek(0)
        
        st.download_button(
            "📊 Download Excel Report",
            excel_buffer,
            "DB_Health_Check_Report.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
    
    with col2:
        # Word Report
        with st.spinner("Generating Word..."):
            word_buffer = generate_word_report(
                results, db_info, st.session_state.summary_inputs,
                st.session_state.frag_thresh, st.session_state.ts_thresh,
                st.session_state.custom_entries, section_order,
                st.session_state.primary_color
            )
        word_buffer.seek(0)
        
        st.download_button(
            "📄 Download Word Report",
            word_buffer,
            "DB_Health_Check_Report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col3:
        # HTML Report
        
        with st.spinner("Generating HTML..."):
            html_report = generate_enhanced_html_report(
                results, db_info, st.session_state.summary_inputs,
                st.session_state.frag_thresh, st.session_state.ts_thresh,
                st.session_state.custom_entries, section_order,
                st.session_state.primary_color
            )
        
        st.download_button(
            "🌐 Download HTML Report",
            html_report.encode('utf-8'),
            "DB_Health_Check_Report.html",
            "text/html",
            use_container_width=True
        )
    
    st.markdown("---")
    st.success("✅ All reports ready for download!")

# UPDATED REPLACEMENT CODE FOR CHECKUP.PY (WITHOUT HELP SECTION & SECURITY NOTE)
# Replace lines 3405-3420 with this cleaner version

else:
    # No data available - show enhanced professional welcome screen
    
    # Add custom CSS for professional styling
    st.markdown("""
    <style>
    .report-option-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 30px;
        margin: 15px 0;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        transition: transform 0.3s ease, box-shadow 0.3s ease;
    }
    
    .report-option-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 15px 40px rgba(0,0,0,0.3);
    }
    
    .option-header {
        color: white;
        font-size: 24px;
        font-weight: bold;
        margin-bottom: 15px;
        display: flex;
        align-items: center;
        gap: 12px;
    }
    
    .option-icon {
        font-size: 32px;
        background: rgba(255,255,255,0.2);
        padding: 10px 15px;
        border-radius: 10px;
        backdrop-filter: blur(10px);
    }
    
    .option-description {
        color: rgba(255,255,255,0.95);
        font-size: 16px;
        line-height: 1.6;
        margin: 15px 0;
    }
    
    .feature-list {
        list-style: none;
        padding: 0;
        margin: 15px 0;
    }
    
    .feature-item {
        color: rgba(255,255,255,0.9);
        padding: 8px 0;
        font-size: 15px;
        display: flex;
        align-items: center;
        gap: 10px;
    }
    
    .feature-item:before {
        content: "✓";
        background: rgba(255,255,255,0.3);
        color: white;
        width: 24px;
        height: 24px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: bold;
        flex-shrink: 0;
    }
    
    .welcome-header {
        text-align: center;
        padding: 30px 0;
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 15px;
        margin-bottom: 30px;
    }
    
    .welcome-title {
        font-size: 36px;
        font-weight: bold;
        color: #2d3748;
        margin-bottom: 10px;
    }
    
    .welcome-subtitle {
        font-size: 18px;
        color: #4a5568;
    }
    
    .divider {
        margin: 30px 0;
        text-align: center;
        position: relative;
    }
    
    .divider:before {
        content: "";
        position: absolute;
        top: 50%;
        left: 0;
        right: 0;
        height: 1px;
        background: linear-gradient(to right, transparent, #cbd5e0, transparent);
    }
    
    .divider-text {
        background: white;
        padding: 0 20px;
        position: relative;
        color: #718096;
        font-weight: 600;
        font-size: 14px;
        letter-spacing: 1px;
    }
    
    .card-gradient-1 {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    .card-gradient-2 {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Welcome Header
    st.markdown("""
    <div class="welcome-header">
        <div class="welcome-title">🏥 Oracle Database Health Check Dashboard</div>
        <div class="welcome-subtitle">Choose your preferred method to generate comprehensive health reports</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Create two columns for the options
    col1, col2 = st.columns(2, gap="large")
    
    with col1:
        st.markdown("""
        <div class="report-option-card card-gradient-1">
            <div class="option-header">
                <span class="option-icon">📄</span>
                <span>Option 1: Upload HTML Report</span>
            </div>
            <div class="option-description">
                Quick analysis of existing health check reports. Perfect for reviewing previously generated database reports.
            </div>
            <ul class="feature-list">
                <li class="feature-item">Upload existing HTML health check report</li>
                <li class="feature-item">Instant analysis and visualization</li>
                <li class="feature-item">Customize sections and insights</li>
                <li class="feature-item">Export to multiple formats (PDF, DOCX, Excel)</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="report-option-card card-gradient-2">
            <div class="option-header">
                <span class="option-icon">🔌</span>
                <span>Option 2: Connect to Database</span>
            </div>
            <div class="option-description">
                Real-time health check execution directly on your Oracle database. Get live insights and immediate results.
            </div>
            <ul class="feature-list">
                <li class="feature-item">Upload CHECKUP_NEW.txt script</li>
                <li class="feature-item">Secure Oracle database connection</li>
                <li class="feature-item">Execute health check in real-time</li>
                <li class="feature-item">Live monitoring and instant dashboard</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
    
    # Divider - Simple and clean
    st.markdown("""
    <div class="divider">
        <span class="divider-text">GET STARTED ABOVE ▲</span>
    </div>
    """, unsafe_allow_html=True)